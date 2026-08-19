"""
BOM(Excel) 与 PDF 原理图 逐器件核对工具
========================================
功能：
  1. 读取excel的任意列
  2. 精确解码 PDF 原理图文字（原理图的字体使用"字形编号+29=ASCII码"的自定义编码，
     需从内容流中还原，OCR 无法保证完整识别）
  3. 二者一一核对：一致(绿)/不一致(红) 全部列出
  4. 生成 Excel 报告(.xlsx) 与文本报告(.txt)

    author : 乔金成
    Time : 2026-08-19
适用：PAT - 本项目仅针对该系列 PDF(原理图导出格式)。
"""

import os  
import re   #正则表达式，用于字符串匹配，查找，替换，文本提取
import sys  # python 解释器，命令行参数，脚本控制，修改搜索路径
import time 
import json  # 保存/恢复上次配置
import math  # 数值判断
import io    # 处理文本流编码
import csv   # 读取 CSV/TSV/文本表格
import threading # 多线程，开启多个字现成并发执行IO，
import traceback   # 大于异常堆栈和模块
from collections import defaultdict #把`collections`里面的`defaultdict`拿过来直接用

# ---------- 依赖检查 ----------
REQUIRED = []
try:
    import pikepdf
except ImportError:
    REQUIRED.append("pikepdf")
try:
    import pdfplumber
except ImportError:
    REQUIRED.append("pdfplumber")
try:
    import openpyxl
except ImportError:
    REQUIRED.append("openpyxl")

if REQUIRED or True:
    import pikepdf  # 再次导入确保错误清晰
    import pdfplumber
    import openpyxl

# ---------- PDF 解码 -------------------------------
DESIGNATOR_RE = re.compile(r"^[A-Za-z]{1,4}\d{1,4}(?:[A-Za-z]\d{1,2})?$")
PREFIXES_PARTS = set("CRLUDQJRTFXSWYKT")


def _build_code2glyph(fonts):
    """把 Type3 字体的 Differences 数组解析成 {字符码: 字形名}。"""
    code2glyph = {} # 定义字典，
    for fn in fonts.keys(): # fn 为字体字符串，fonts[fn] 为字体对象  
        f = fonts[fn]   # 赋值
        enc = f.get("/Encoding") # 获取字体编码对象
        diffs = None #数组为空
        if enc is not None and hasattr(enc, "get"):  #
            diffs = enc.get("/Differences", None)
        if diffs is None:
            continue
        try:  # 将diffs中的每个元素转换为字符串，并存储在arr列表中
            arr = [str(x) for x in diffs]  # 数组转换
        except Exception: 
            continue
        gmap = {} # 定义字典 gmap
        start = None
        names = [] # 列表 names 名称

        def flush():  # 定义函数
            nonlocal gmap, start, names #循环，将start和names中的元素添加到gmap字典中
            if start is None:
                return
            cur = start
            for g in names:
                gmap[cur] = g
                cur += 1
            start = None  # 处理完成，进行重置
            names = []

        for x in arr: 
            if re.fullmatch(r"-?\d+", x): # 正则表达式，判断元素是否整数（正负数字字符串）
                flush() #调用函数
                start = int(x) 
            else:
                names.append(x)
        flush() #调用，避免数组
        code2glyph[str(fn)] = gmap
    return code2glyph


def decode_page(pdf, pageno): 
    """完整解码一页：返回[(字符, x, y, 字号)]，自原点为页面左上，y 越上越小。"""
    from pikepdf import parse_content_stream #库pikepdf中的函数parse_content_stream，解析PDF内容流

    page = pdf.pages[pageno]  # 获取指定页码的页面对象
    c2g = _build_code2glyph(page["/Resources"]["/Font"])  #`c2g` = code‑to‑glyph：**字体名称 → {字节编码：字形名}**
    c2u = _build_code2unicode(page["/Resources"]["/Font"])  # Type0 字体 ToUnicode 映射 {2字节码: unicode}
    # 传入 page 对象而非 page.Contents：pikepdf 会自动合并 Contents 数组（多流页面）
    ops = list(parse_content_stream(page))
    tm = [1, 0, 0, 1, 0, 0]  #状态变量
    tl = 0.0 #原点00
    cur_font = None #当前字体，默认不使用
    cur_size = 1.0  #缩放系数默认1.0
    out = [] #输出列表，存放解析之后提出的文本内容
    for op in ops:
        name = str(op.operator) 
        ops_ = list(op.operands)
        if name == "BT":
            tm = [1, 0, 0, 1, 0, 0]
        elif name == "Tf":
            cur_font = str(ops_[0])
            try:
                cur_size = float(ops_[1])
            except Exception:
                pass
        elif name == "Tm":
            tm = [float(x) for x in ops_[:6]]
        elif name == "Td":
            tx, ty = float(ops_[0]), float(ops_[1])
            a, b, c, d, e, f = tm
            tm = [a, b, c, d, e + tx * a + ty * c, f + tx * b + ty * d]
        elif name == "T*":
            a, b, c, d, e, f = tm
            tm = [a, b, c, d, e, f + tl * d]
        elif name == "TL":
            tl = float(ops_[0])
        elif name in ("Tj", "'"):
            b = bytes(ops_[0]) if hasattr(ops_[0], "__bytes__") else b""
            if cur_font in c2u:
                _append_c2u(out, c2u[cur_font], b, tm, cur_size)
            else:
                gmap = c2g.get(cur_font, {})
                for byte in b:
                    _append_char(out, gmap, byte, tm, cur_size)
        elif name == "TJ":
            for item in ops_[0]:
                if hasattr(item, "__bytes__"):
                    b = bytes(item)
                    if cur_font in c2u:
                        _append_c2u(out, c2u[cur_font], b, tm, cur_size)
                    else:
                        gmap = c2g.get(cur_font, {})
                        for byte in b:
                            _append_char(out, gmap, byte, tm, cur_size)
    return out


def _build_code2unicode(fonts):
    """把 Type0/Type1/TrueType 字体的 /ToUnicode CMap 解析成 {字体名: {'map':{code:unicode}, 'bytes':N}}。
    bytes 表示码点宽度：1 或 2（由 codespacerange 决定）。"""
    out = {}
    for fn in fonts.keys():
        f = fonts[fn]
        tu = f.get("/ToUnicode", None)
        if tu is None:
            continue
        try:
            raw = tu.read_bytes().decode("latin-1", errors="replace")
        except Exception:
            try:
                raw = bytes(tu).decode("latin-1", errors="replace")
            except Exception:
                continue
        m = {}
        # codespacerange：判断单字节还是双字节码点
        nbytes = 2
        cm = re.search(r"begincodespacerange\s*<(?:0{0,2}[0-9A-Fa-f]+)>\s*<([0-9A-Fa-f]+)>", raw)
        if cm:
            hi = cm.group(1)
            nbytes = 1 if len(hi) <= 2 else 2
        # bfchar: <code> <unicode>
        for cm in re.finditer(r"beginbfchar\s*(.*?)\s*endbfchar", raw, re.S):
            for pair in re.findall(r"<([0-9A-Fa-f]+)>\s*<([0-9A-Fa-f]+)>", cm.group(1)):
                code, uni = int(pair[0], 16), _utf16be_to_str(pair[1])
                m[code] = uni
        # bfrange: <lo> <hi> <dst>（连续映射，dst 可为起始值）
        for cm in re.finditer(r"beginbfrange\s*(.*?)\s*endbfrange", raw, re.S):
            block = cm.group(1)
            # 多行连续形式（三行一组）
            for trio in re.finditer(
                    r"<([0-9A-Fa-f]+)>\s*<([0-9A-Fa-f]+)>\s*<([0-9A-Fa-f]+)>", block):
                lo, hi = int(trio.group(1), 16), int(trio.group(2), 16)
                dst = _utf16be_to_str(trio.group(3))
                if len(dst) == 1:
                    for code in range(lo, hi + 1):
                        m[code] = dst
                        dst = chr(ord(dst) + 1)
                else:
                    for code in range(lo, hi + 1):
                        m[code] = dst
        if m:
            out[str(fn)] = {"map": m, "bytes": nbytes}
    return out


def _utf16be_to_str(hexstr):
    """把 <hex> 的 UTF-16BE 转成字符。"""
    h = hexstr.strip()
    try:
        if len(h) == 4:
            return chr(int(h, 16))
        b = bytes.fromhex(h)
        try:
            return b.decode("utf-16-be")
        except Exception:
            pass
        # 允许 1 字节编码
        try:
            return b.decode("latin-1")
        except Exception:
            return ""
    except Exception:
        return ""


def _append_c2u(out, umap_obj, b, tm, cur_size):
    """按码点宽度解码 ToUnicode 字体（1 或 2 字节，由 codespacerange 决定）。"""
    umap, nbytes = umap_obj.get("map", {}), umap_obj.get("bytes", 2)
    i = 0
    n = len(b)
    while i < n:
        code = b[i]
        if nbytes == 2 and i + 1 < n:
            code = (b[i] << 8) | b[i + 1]
        ch = umap.get(code)
        if ch is None and nbytes == 2:
            ch = umap.get(b[i])  # 回退：按单字节查一次
        if ch is None:
            ch = "?"
        for c in ch:
            out.append((c, tm[4], tm[5], cur_size))
        i += nbytes



def _append_char(out, gmap, byte, tm, cur_size):
    glyph = gmap.get(byte)
    m = re.match(r"/g(\d+)$", glyph or "")
    num = int(m.group(1)) if m else None
    ch = chr(num + 29) if num is not None and 0 < num + 29 <= 0x10FFFF else "?"
    out.append((ch, tm[4], tm[5], cur_size))

# 根据传入的字节值（byte）在字形映射表（gmap）中查找对应的字形名称，解析出字形编号，
# 并将其转换为实际的字符，最后将该字符连同坐标与字号信息一起存入输出列表中。

def decode_pdf_blocks(pdf_path):
    """解码 PDF 全部页面，返回页面文字行(列表) 与 全部字符。
    页序 1..N。每行保留位置信息用于后续定位。
    """
    pdf = pikepdf.open(pdf_path)
    pages_words = {}     # page -> {y: [(word, x0, x1)]}  （已按行分词组）
    pages_text = {}      # page -> [行文字]
    page_count = len(pdf.pages)

    for pno in range(page_count):
        chars = decode_page(pdf, pno)
        # 与 pdfplumber 的精确 bbox 对齐（保障字符宽度/间距正确）
        boxes = _get_char_boxes(pdf_path, pno)
        if boxes and len(boxes) == len(chars):
            merged = []
            for (ch, x, y, sz), (x0, top, x1, bottom) in zip(chars, boxes):
                merged.append((ch, y, sz, x0, x1, top))
        else:
            merged = [(ch, y, sz, x, x, y) for (ch, x, y, sz) in chars]

        # 以 pdfplumber 的 top(同行近似)为行键
        lines = defaultdict(list)
        for item in merged: # 遍历合并后的字符信息
            ch, y, sz, x0, x1, top = item # 解包字符信息
            key = round(top, 7) # 将 top 坐标四舍五入到小数点后 7 位，作为行的键
            lines[key].append((ch, x0, x1, sz, top)) # 将字符信息添加到对应行的列表

        page_words = []  # 每页的行列表，每行包含字符及其位置信息
        page_text = []  # 每页的行文字列表，每行仅包含文本内容
        for yk in sorted(lines.keys()): # 按行键排序，确保行的顺序与 PDF 中的顺序一致
            seq = sorted(lines[yk], key=lambda t: t[1]) # 按字符的 x0 坐标排序，确保字符在行内的顺序正确
            words = [] # 存储当前行的单词列表，每个单词包含文本及其位置信息
            cur = "" # 当前正在构建的单词文本
            cx0 = None # 当前单词的起始 x 坐标
            prev_x1 = None # 上一个字符的结束 x 坐标
            line_text = ""  # 当前行的完整文本内容
            gap_sz = None   # 当前字符间距阈值，用于判断是否需要分词
            for ch, x0, x1, sz, top in seq: # 遍历当前行的字符信息
                line_text += ch # 将字符添加到当前行的完整文本内容中
                if cur == "": # 如果当前单词为空，说明这是新单词的开始
                    cur = ch # 将当前字符作为新单词的起始字符
                    cx0 = x0
                    gap_sz = max(0.5, sz * 0.30)
                else:
                    if x0 - prev_x1 > gap_sz:
                        words.append((cur, cx0, prev_x1, top))
                        cur = ch
                        cx0 = x0
                    else:
                        cur += ch
                    gap_sz = max(0.5, sz * 0.30)
                prev_x1 = x1
            if cur:
                words.append((cur, cx0, prev_x1, top))
            page_words.append((yk, words))
            page_text.append((yk, line_text))
        pages_words[pno + 1] = page_words
        pages_text[pno + 1] = page_text
        pages_words[pno + 1] = page_words
        pages_text[pno + 1] = page_text
    pdf.close()
    return pages_words, pages_text, page_count


def _get_char_boxes(pdf_path, pageno):
    try:
        with pdfplumber.open(pdf_path) as p:
            return [
                (c["x0"], c["top"], c["x1"], c["bottom"])
                for c in p.pages[pageno].chars
            ]
    except Exception:
        return None


def _split_designator_word(up):
    """把一个解码词拆成若干位号（处理 'C2, C7'、'C300-C302'、'R2180R0201' 等紧凑写法）。"""
    parts = re.split(r"[,;\s\u3001]+", up)
    out = []
    for part in parts:
        if DESIGNATOR_RE.match(part) and part[0] in PREFIXES_PARTS:
            out.append(part)
            continue
        # 范围展开：C300-C302 / R1-R5
        m = re.fullmatch(r"([A-Za-z]{1,4})(\d{1,4})-([A-Za-z]{1,4})(\d{1,4})", part)
        if m:
            p1, n1s, p2, n2s = m.groups()
            if p1 == p2:
                try:
                    n1, n2 = int(n1s), int(n2s)
                    if 0 <= n1 <= n2 <= 99999 and n2 - n1 <= 200:
                        out += [p1 + str(n) for n in range(n1, n2 + 1)]
                        continue
                except Exception:
                    pass
        # 紧凑拼接：从词中提取所有位号，如 "HP186R218" → R218、"R2180R0201" → R218
        # 用连续片段扫描 [字母1-4][数字1-4]，且首个字母属于器件前缀集合
        run = re.match(r"([A-Za-z]{1,4}\d{1,4})*[A-Za-z]{1,4}\d{1,4}", part)
        if run:
            seg = run.group(0)
            for fm in re.finditer(r"([A-Za-z]{1,4})(\d{1,4})", seg, re.I):
                pfx, num = fm.group(1), fm.group(2)
                if pfx.upper()[0] in PREFIXES_PARTS and DESIGNATOR_RE.match(pfx + num):
                    out.append((pfx + num).upper())
    return out


def extract_designators(pages_words):
    """从解码文字中提取器件位号：{位号: [页码,...]}（支持紧凑/逗号/范围/跨行范围）。"""
    tokens = defaultdict(list)
    all_words = []
    for pno, pwords in pages_words.items():
        open_range = None  # 上一词末尾未闭合的范围前缀，如 ('R','86')，来自 "R86-"
        for y, words in pwords:
            for w, x0, x1, top in words:
                up = w.upper().strip()
                if not up:
                    continue  # 空词不打断跨行范围
                all_words.append((pno, y, up, x0, x1, top))

                merged = False
                # 跨行范围：上一词 "R86-" → 本词 "R89" / "89," 
                if open_range:
                    pfix, n1 = open_range
                    mo = re.match(r"((?:[A-Za-z]{1,4})?)(\d{1,4})", up)
                    n2 = None
                    if mo:
                        n2 = int(mo.group(2))
                        mergable = mo.group(1) == "" or mo.group(1) == pfix
                    else:
                        mergable = False
                    if n2 is not None and mergable and 0 <= n1 <= n2 <= 99999 and n2 - n1 <= 200:
                        for n in range(n1, n2 + 1):
                            tokens[pfix + str(n)].append(pno)
                        merged = True
                    open_range = None

                if merged:
                    continue

                for d in _split_designator_word(up):
                    tokens[d].append(pno)

                # 若本词以 "Rxx-" 结尾，保留供下一词合并（如 "R86-")
                mm = re.search(r"([A-Za-z]{1,4}\d{1,4})-+\s*$", up)
                if mm:
                    base = re.fullmatch(r"([A-Za-z]{1,4})(\d+)", mm.group(1))
                    if base:
                        open_range = (base.group(1), int(base.group(2)))
    return tokens, all_words


def build_pdf_designator_annotations(pages_words):
    """给 PDF 里每个位号找邻近标注词（值/封装/型号）。
    返回 {位号: {'pages':[页码], 'near':[附近词,去重]}}。
    """
    tokens, _ = extract_designators(pages_words)
    # 每页的所有词
    page_words = {} # 
    for pno, pwords in pages_words.items():
        plist = []
        for y, words in pwords:
            for w, x0, x1, top in words:
                if re.match(r"^[\w.+/@-]{2,}$", w):
                    plist.append({"y": y, "x": x0, "w": w.upper()})
        page_words[pno] = plist

    out = {}
    for des, pages in tokens.items():
        des_pos = None
        for pno, pwords in pages_words.items():
            if pno not in pages:
                continue
            for y, words in pwords:
                for w, x0, x1, top in words:
                    if w.upper() == des:
                        des_pos = (pno, y, x0)
                        break
                if des_pos:
                    break
            if des_pos:
                break
        near = []
        if des_pos:
            pno, dy, dx = des_pos
            for it in page_words.get(pno, []):
                if abs(it["y"] - dy) <= 120 and abs(it["x"] - dx) <= 260 and it["w"] != des:
                    near.append(it["w"])
        out[des] = {"pages": pages, "near": sorted(set(near))[:15]}
    return out


def app_smart(v):
    """模块级包装，避免重复定义 smart_norm_value。"""
    try:
        return smart_norm_value(v)
    except Exception:
        return v


def compare_pdf_to_excel(pdf_path, bom_path, primary="Part Reference"):
    """以 PDF 位号为准，检查 Excel 中是否存在且 Value 是否疑似一致。
    返回 rows 与 stats。"""
    pages_words, pages_text, _ = decode_pdf_blocks(pdf_path)
    annotations = build_pdf_designator_annotations(pages_words)

    headers, data = load_bom(bom_path)
    primary = primary if primary in headers else (_designator_col_name(headers) or headers[0])
    pidx = headers.index(primary)
    vidx = headers.index("Value") if "Value" in headers else None
    fidx = headers.index("PCB Footprint") if "PCB Footprint" in headers else None
    qidx = headers.index("Quantity") if "Quantity" in headers else None

    excel_index = {}
    for row in data:
        for d in split_designators_text(row["values"][pidx]):
            excel_index.setdefault(d.upper(), []).append(row)

    rows = []
    n_found = n_pdf_only = n_ok = n_may = 0
    for des in sorted(annotations.keys()):
        ann = annotations[des]
        near = ann["near"]
        near_text = "; ".join(near[:8])
        excel_rows = excel_index.get(des)
        if not excel_rows:
            n_pdf_only += 1
            rows.append({"item": des, "status": "PDF有Excel无", "valueA": "", "valueB": "",
                         "footA": "", "footB": "", "near": near_text, "qty": ""})
            continue
        n_found += 1
        er = excel_rows[0]
        ev = er["values"][vidx] if vidx is not None else ""
        ef = er["values"][fidx] if fidx is not None else ""
        eq = er["values"][qidx] if qidx is not None else ""
        # 值/封装是否能在邻近标注中找到（直接命中 / 归一化命中 / 短数值允许子串命中）
        nv = app_smart(ev)
        nf = app_smart(ef)
        hit = False
        candidates = [str(x) for x in [ev, nv, ef, nf] if x]
        for c in candidates:
            c = str(c).strip().lower()
            if not c:
                continue
            for w in near:
                wl = w.lower()
                if c == wl:
                    hit = True
                    break
                if app_smart(c) == app_smart(w):
                    hit = True
                    break
                # 仅当候选词是短数值(<=5字符)时才做子串匹配(值如 49.9/10K 常用)
                # 型号/封装等长词只整词匹配，避免 "13" 误命中 "TMUX1308..."
                c_short = re.fullmatch(r"[\d.]{1,6}", c)
                if c_short and len(c) <= 5 and (c in wl or wl in c):
                    hit = True
                    break
            if hit:
                break
        if hit:
            n_ok += 1
            rows.append({"item": des, "status": "一致", "valueA": ev, "valueB": "",
                         "footA": ef, "footB": "", "near": near_text, "qty": eq})
        else:
            n_may += 1
            rows.append({"item": des, "status": "待确认(值/封装在PDF上未匹配)", "valueA": ev,
                         "valueB": "", "footA": ef, "footB": "", "near": near_text, "qty": eq})

    stats = {
        "all": {
            "主键": primary,
            "PDF位号总数": len(annotations),
            "Excel中找到": n_found,
            "PDF有Excel无": n_pdf_only,
            "值疑似一致": n_ok,
            "值疑似不一致/待确认": n_may,
        },
        "bad": [{"designator": r["item"], "row": ""} for r in rows
                if r["status"] != "一致"],
        "extra": [r["item"] for r in rows if r["status"] == "PDF有Excel无"],
        "in_bom": {r["item"] for r in rows if r["status"] != "PDF有Excel无"},
        "common": [r["item"] for r in rows if r["status"] == "一致"],
        "only_a": [r["item"] for r in rows if r["status"] == "PDF有Excel无"],
        "only_b": [],
        "mode": "pdf2excel",
    }
    return rows, stats


# ---------- MPN 对比：以 MPN 为查询主键 --------------------------------
_MPN_COL_KEYS = ("mpn", "manufacturer part number", "manufacturer pn", "mfr pn",
                 "manufacturer item number", "manufacturer item", "part number", "part no",
                 "料号", "型号")
_MFG_COL_KEYS = ("manufacturer", "manufacture", "mfg", "vendor", "厂家", "供应商", "厂商")
_QTY_COL_KEYS = ("quantity", "qty", "数量", "pcs")
_REF_COL_KEYS = ("reference", "refdes", "designator", "refs", "位号", "ref", "references",
                 "reference designator")
_ITM_COL_KEYS = ("item number", "part number", "part no", "item no", "ipn", "零件号", "料号")
_DESC_COL_KEYS = ("item name", "description", "desc", "part", "value", "名称", "描述", "spec")


def _find_col(headers, keys):
    """在表头中查找指定语义的列，返回列名或 None。"""
    low = [str(x or "").lower().replace(" ", "") for x in headers]
    for k in keys:
        kl = k.lower().replace(" ", "")
        for i, n in enumerate(low):
            if n == kl or n.endswith(kl) or n.startswith(kl):
                return headers[i]
    return None


def _extract_pdf_mpn_index(pages_words):
    """从 PDF 解码文字中提取 {MPN: {pages:[...], near_refs:set, qty}}。
    MPN 判定：6 个以上字符，含字母与数字（或含 - 的长料号）。"""
    mpn_re = re.compile(r"^(?=.*[A-Za-z])(?=.*\d)[A-Za-z0-9][A-Za-z0-9.\-]{4,}$")
    out = defaultdict(list)
    # 先收集每页每个词及其位置，再找 MPN 附近的位号
    for pno, pwords in pages_words.items():
        for y, words in pwords:
            for w, x0, x1, top in words:
                up = w.upper().strip()
                if not up:
                    continue
                if mpn_re.match(up) and up[0].isalnum() and len(re.sub(r"[\-.0-9]", "", up)) >= 1:
                    # 扩展范围：R2180R0201 这类粘连词里也拆出 MPN
                    for cand in _mpn_candidates(up):
                        out[cand].append((pno, y, x0))
    # 去重 page
    index = {}
    for mpn, hits in out.items():
        uniq = []
        seen = set()
        for p, y, x in hits:
            if (p, x) not in seen:
                uniq.append((p, y, x))
                seen.add((p, x))
        index[mpn] = uniq
    return index


def _mpn_candidates(up):
    """从一个词中切出可能的 MPN（优先整词；粘连词按 连续[字母+数字+-] 片段切）。"""
    if re.fullmatch(r"[A-Za-z0-9][A-Za-z0-9.\-]{4,}", up) and any(c.isdigit() for c in up):
        return [up]
    cands = []
    for m in re.finditer(r"[A-Za-z][A-Za-z0-9\-]{4,}\d[A-Za-z0-9\-]*", up):
        c = m.group(0)
        if any(cd in c for cd in "0123456789") and len(c) >= 5:
            cands.append(c)
    return cands


def _linked_designators(pages_words, pno, y, x, mpn, radius=90):
    """找同一页离 MPN 最近的位号词（排除 MPN 自身的碎片词）。
    返回按距离排序的去重位号列表。"""
    mpn_clean = re.sub(r"[^A-Za-z0-9]", "", (mpn or "").upper())
    cands = []
    for yy, words in pages_words.get(pno, []):
        for w, x0, x1, top in words:
            up = w.upper().strip()
            for d in _split_designator_word(up):
                if not DESIGNATOR_RE.match(d) or d[0] not in PREFIXES_PARTS or len(d) > 6:
                    continue
                # 排除 MPN 自身的碎片（如 TCAN1146、LVC1、DRV8316 等）
                if mpn_clean and (mpn_clean == d or mpn_clean.startswith(d) and len(d) >= 4):
                    continue
                dm = max(abs(yy - y), abs(x0 - x))
                if dm <= radius:
                    cands.append((dm, d))
    cands.sort()
    seen = set()
    out = []
    for _dm, d in cands:
        if d not in seen:
            seen.add(d)
            out.append(d)
    return out


def compare_excel_pdf_by_mpn(bom_path, pdf_path, mpn_col=None):
    """以 MPN(料号) 为查询主键，核对 Excel BOM 与 PDF 原理图是否一致。
    一致时进一步比对 数量(Quantity)、位号(Reference)、厂商(Manufacturer)。
    返回 rows 与 stats。"""
    headers, data = load_bom(bom_path)
    # 仅当传入列名在表头中存在时才采用，否则自动识别 MPN/料号 列
    if mpn_col in headers:
        pass
    else:
        mpn_col = None
    mpn_col = mpn_col or _find_col(headers, _MPN_COL_KEYS) or headers[0]
    qty_col = _find_col(headers, _QTY_COL_KEYS)
    ref_col = _find_col(headers, _REF_COL_KEYS)
    mfg_col = _find_col(headers, _MFG_COL_KEYS)
    if mpn_col not in headers:
        raise RuntimeError("找不到 MPN/料号 列")
    mi = headers.index(mpn_col)
    qi = headers.index(qty_col) if qty_col in headers else None
    ri = headers.index(ref_col) if ref_col in headers else None
    fi = headers.index(mfg_col) if mfg_col in headers else None

    pages_words, pages_text, _ = decode_pdf_blocks(pdf_path)
    pdf_mpn = _extract_pdf_mpn_index(pages_words)

    # Excel 主键 -> 记录
    excel_rows = []
    for row in data:
        mpn = str(row["values"][mi] or "").strip().upper()
        if not mpn:
            continue
        excel_rows.append({
            "mpn": mpn,
            "mpn_raw": row["values"][mi],
            "qty": ("" if qi is None else row["values"][qi]) or "",
            "qty_n": (str(row["values"][qi]).strip() if qi is not None and row["values"][qi] not in (None, "") else ""),
            "refs": split_designators_text(row["values"][ri]) if ri is not None else [],
            "mfg": ("" if fi is None else row["values"][fi]) or "",
        })

    rows = []
    seens = set()
    for er in excel_rows:
        mpn = er["mpn"]
        if mpn in seens:
            continue
        seens.add(mpn)
        phits = pdf_mpn.get(mpn, [])
        pages = sorted({p for p, _, _ in phits})
        if not phits:
            rows.append({
                "item": mpn, "mpn": mpn, "status": "仅Excel有(PDF无此料号)",
                "qty_a": er["qty_n"], "qty_b": "", "refs_a": ",".join(er["refs"]),
                "refs_b": "", "mfg": str(er["mfg"]), "pages": "", "mismatchs": ["PDF中未找到该MPN"],
            })
            continue
        # PDF 就近位号（尽力而为：每个命中点取最近位号，跨页去重）
        pdf_refs = []
        seen_refs = set()
        for p, y, x in phits:
            for d in _linked_designators(pages_words, p, y, x, mpn):
                if d not in seen_refs:
                    seen_refs.add(d)
                    pdf_refs.append(d)
        # 提示信息（不判死，仅提示）
        notes = []
        eqty = None
        try:
            eqty = int(float(str(er["qty_n"]).replace(",", "").strip()))
        except Exception:
            eqty = None
        # 跨页大符号只按首页就近位号估算
        p0, y0, x0 = phits[0]
        pdf_est = _linked_designators(pages_words, p0, y0, x0, mpn)
        if eqty is not None and pdf_est and eqty != len(pdf_est):
            notes.append("Excel数量=%s 但PDF就近位号≈%s" % (eqty, len(pdf_est)))
        er_set = set(er["refs"])
        missing = sorted(er_set - set(pdf_est))
        if missing and pdf_est:
            notes.append("PDF就近未到位号: %s" % ",".join(missing[:10]))
        # 状态：以 MPN 是否在 PDF 中找到为主，差异作为待确认提示
        st = "一致" if not notes else "待确认(" + "; ".join(notes[:2]) + ")"
        rows.append({
            "item": mpn, "mpn": mpn, "status": st,
            "qty_a": er["qty_n"], "qty_b": str(len(pdf_refs)),
            "refs_a": ",".join(sorted(er["refs"])),
            "refs_b": ",".join(sorted(pdf_refs)),
            "mfg": str(er["mfg"]), "pages": ",".join(str(p) for p in pages),
            "mismatchs": notes,
        })

    pdf_only = sorted(set(pdf_mpn.keys()) - seens)
    n_ok = sum(1 for r in rows if r["status"] == "一致")
    n_confirm = sum(1 for r in rows if r["status"].startswith("待确认"))
    n_only_a = sum(1 for r in rows if r["status"].startswith("仅Excel有"))
    stats = {
        "all": {
            "主键": mpn_col,
            "Excel 料号(MPN)总数": len(seens),
            "PDF 识别料号总数": len(pdf_mpn),
            "一致(数量/位号匹配)": n_ok,
            "待确认(数量或位号提示)": n_confirm,
            "仅Excel有(PDF无此料号)": n_only_a,
            "PDF有而Excel无": len(pdf_only),
        },
        "bad": [{"designator": r["item"], "row": ""} for r in rows if r["status"] != "一致"],
        "extra": pdf_only,
        "in_bom": {r["item"] for r in rows},
        "common": [r["item"] for r in rows if r["status"] == "一致"],
        "only_a": [r["item"] for r in rows if r["status"].startswith("仅Excel有")],
        "only_b": pdf_only,
        "mode": "mpn",
    }
    return rows, stats


# ---------- BOM 读取 --------------------------------
_CSV_ENCODINGS = ["utf-8-sig", "utf-8", "gb18030", "gbk", "big5", "utf-16", "latin-1"]

# 常见 BOM 表头关键词（用于自动定位表头行 / 位号列）
_DESIGNATOR_KEYS = ("part reference", "reference", "refdes", "ref des", "designator",
                    "位号", "refs", "references", "device designator", "comp designator")
_HEADER_KEYS = ("reference", "value", "quantity", "qty", "footprint", "package",
                "mpn", "part number", "part no", "manufacturer", "description",
                "item", "designator", "位号", "型号", "数量", "封装", "厂家", "料号", "描述")


def _find_header_row(rows):
    """在多行表格中自动定位表头行（有些 BOM 第 4 行才是表头）。
    返回 (表头行在 rows 中的下标, 是否成功)。"""
    best_i, best_score, best_has_des = -1, -1, False
    for i, row in enumerate(rows[:24]):  # 只在前 24 行找
        cells = [re.sub(r"\s+", "", str(c or "")).lower() for c in row]
        score = 0
        has_des = False
        joined = " ".join(cells)
        for k in _HEADER_KEYS:
            if k in joined:
                score += 1
        for k in _DESIGNATOR_KEYS:
            if k in joined and any(c.startswith(k.replace(" ", "")) or
                                   k.replace(" ", "") in c for c in cells):
                has_des = True
                score += 2
        if has_des and score > best_score:
            best_i, best_score, best_has_des = i, score, True
    if best_i < 0 and best_score < 2:
        return 0, False
    return max(best_i, 0), True


def _detect_designator_col(headers):
    """在表头列表中智能识别位号列（支持 Reference / Part Reference / 位号…）。"""
    names = [str(x or "").strip() for x in headers]
    low = [n.lower().replace(" ", "") for n in names]
    for k in _DESIGNATOR_KEYS:
        kl = k.replace(" ", "")
        for i, n in enumerate(low):
            if n == kl or n.endswith(kl) or n.startswith(kl):
                return i
    return 0


def _designator_col_name(headers):
    """返回位号列名（未识别时返回第 1 列名）。"""
    if not headers:
        return None
    return str(headers[_detect_designator_col(headers)])


def _csv_content_rows(bom_path):
    """读取 CSV/TSV/文本表格，自动探测编码与分隔符，返回原始行(字符串列表)。"""
    raw = open(bom_path, "rb").read()   # 读成字节流，探测编码
    enc = None
    text = None
    for cand in _CSV_ENCODINGS:
        try:
            text = raw.decode(cand)
            enc = cand
            break
        except Exception:
            continue
    if text is None:
        text = raw.decode("latin-1", errors="replace")
        enc = "latin-1"
    sample = text[:4096]
    try:
        dialect = csv.Sniffer().sniff(sample, delimiters=",;\t| ")
        delim = dialect.delimiter
    except Exception:
        delim = ","  # 探测不到就默认逗号
    return [row for row in csv.reader(io.StringIO(text), delimiter=delim)], enc


def _load_xls(bom_path):
    """读取旧版 .xls(Excel 97-2003)，依赖 xlrd。"""
    try:
        import xlrd
    except ImportError:
        raise RuntimeError("读取 .xls 需要安装 xlrd：pip install xlrd")
    # xlrd>=2.0 只支持 .xls，不支持 .xlsx，这里仅处理 .xls
    wb = xlrd.open_workbook(bom_path)
    ws = wb.sheet_by_index(0)
    headers = []
    data = []
    for r in range(ws.nrows):
        row = [ws.cell_value(r, c) for c in range(ws.ncols)]
        row = ["" if v is None else v for v in row]
        if r == 0:
            headers = [str(x) for x in row]
        else:
            data.append({"header_row": r + 1, "values": row})
    return headers, data


def _load_xlsb(bom_path):
    """读取 Excel 二进制 .xlsb，依赖 pyxlsb。"""
    try:
        import pyxlsb
    except ImportError:
        raise RuntimeError("读取 .xlsb 需要安装 pyxlsb：pip install pyxlsb")
    headers = []
    data = []
    with pyxlsb.open_workbook(bom_path) as wb:
        ws = wb.get_sheet_by_name(wb.sheetnames[0])
        r = 0
        for row in ws.rows():
            r += 1
            vals = [("" if c is None else c.v) for c in row]
            if r == 1:
                headers = [str(x) for x in vals]
            else:
                data.append({"header_row": r, "values": vals})
    return headers, data


def _load_ods(bom_path):
    """读取 LibreOffice/OpenOffice .ods，依赖 odfpy。"""
    try:
        from odf.opendocument import load
        from odf.table import Table, TableRow, TableCell
        from odf.text import P
    except ImportError:
        raise RuntimeError("读取 .ods 需要安装 odfpy：pip install odfpy")
    doc = load(bom_path)
    headers = []
    data = []
    for tbl in doc.spreadsheet.getElementsByType(Table):
        r = 0
        for row in tbl.getElementsByType(TableRow):
            r += 1
            vals = []
            for cell in row.getElementsByType(TableCell):
                txt = ""
                for p in cell.getElementsByType(P):
                    for node in p.childNodes:
                        if node.nodeType == node.TEXT_NODE:
                            txt += node.data
                        elif node.nodeType == node.ELEMENT_NODE and node.qName == "a":
                            txt += "".join(n.data for n in node.childNodes if n.nodeType == n.TEXT_NODE)
                # 处理单元格重复列（number-columns-repeated）
                repeat = 1
                try:
                    repeat = int(cell.getAttribute("numbercolumnsrepeated") or 1)
                except Exception:
                    pass
                for _ in range(repeat):
                    vals.append(txt)
            if r == 1:
                headers = [str(x) for x in vals]
            else:
                data.append({"header_row": r, "values": vals})
        break  # 只取第一个 sheet
    if not headers:
        raise RuntimeError("ODS 文件未找到表格数据")
    return headers, data


def _load_html(bom_path):
    """读取 HTML 表格(.html/.htm)中第一个 <table>，依赖内置模块。"""
    import html.parser

    class _TableParser(html.parser.HTMLParser):
        def __init__(self):
            super().__init__()
            self.in_table = False
            self.in_cell = False
            self.in_row = False
            self.cur_cell = []
            self.cur_row = []
            self.tables = []

        def handle_starttag(self, tag, attrs):
            t = tag.lower()
            if t == "table":
                self.in_table = True
                self.tables.append([])
            elif t == "tr" and self.in_table:
                self.cur_row = []
            elif t in ("td", "th") and self.in_table:
                self.in_cell = True
                self.cur_cell = []
            elif t in ("br", "p") and self.in_cell:
                self.cur_cell.append(" ")

        def handle_data(self, data):
            if self.in_cell:
                self.cur_cell.append(data)

        def handle_endtag(self, tag):
            t = tag.lower()
            if t in ("td", "th") and self.in_cell:
                self.cur_row.append("".join(self.cur_cell).strip())
                self.in_cell = False
            elif t == "tr":
                if self.cur_row:
                    self.tables[-1].append(self.cur_row)
            elif t == "table":
                self.in_table = False

    with open(bom_path, "rb") as f:
        raw = f.read()
    text = None
    for cand in _CSV_ENCODINGS:
        try:
            text = raw.decode(cand)
            break
        except Exception:
            continue
    parser = _TableParser()
    parser.feed(text or raw.decode("latin-1", errors="replace"))
    if not parser.tables:
        raise RuntimeError("HTML 文件中未找到 <table>")
    table = parser.tables[0]
    if not table:
        raise RuntimeError("HTML 表格为空")
    headers = [str(x) for x in table[0]]
    data = [{"header_row": i + 1, "values": row} for i, row in enumerate(table[1:])]
    return headers, data


def load_bom(bom_path, sheet_contains=None):
    """读取表格文件(xlsx/csv/tsv/txt/xls/xlsb/ods/html...)，返回列名列表与行数据(索引从1开始，含表头行号)。

    支持格式：
      - Excel(.xlsx/.xlsm)        —— openpyxl (内置)
      - 旧版 Excel(.xls)          —— 需 pip install xlrd
      - Excel 二进制(.xlsb)       —— 需 pip install pyxlsb
      - LibreOffice(.ods)         —— 需 pip install odfpy
      - HTML 表格(.html/.htm)     —— 内置解析
      - CSV/TSV/文本表格          —— 内置 csv 模块，自动探测编码(UTF-8/GBK/…)与分隔符
    """
    ext = os.path.splitext(str(bom_path).lower())[1]
    if ext == ".xls":
        return _load_xls(bom_path)
    if ext == ".xlsb":
        return _load_xlsb(bom_path)
    if ext == ".ods":
        return _load_ods(bom_path)
    if ext in (".html", ".htm"):
        return _load_html(bom_path)
    if ext in (".csv", ".tsv", ".txt", ".tab", ".text"):
        rows, enc = _csv_content_rows(bom_path)
        rows = [[str(v) if not (v is None or isinstance(v, float) and math.isnan(v)) else ""
                 for v in row] for row in rows]
        hi, _ = _find_header_row(rows)
        headers = [str(x) for x in rows[hi]]
        data = [{"header_row": i + 1, "values": row}
                for i, row in enumerate(rows[hi + 1:], start=hi + 1) if any(str(x).strip() for x in row)]
        return headers, data
    # 默认按 xlsx 读
    wb = openpyxl.load_workbook(bom_path, data_only=True)
    ws = wb.active
    allrows = []
    for r in range(1, ws.max_row + 1):
        row = []
        for c in range(1, ws.max_column + 1):
            v = ws.cell(row=r, column=c).value
            row.append(v if v is not None else "")
        allrows.append(row)
    hi, _ = _find_header_row(allrows)
    headers = [str(x) for x in allrows[hi]]
    data = [{"header_row": r, "values": row}
            for r, row in enumerate(allrows[hi + 1:], start=hi + 2) if any(str(x).strip() for x in row)]
    return headers, data


def split_designators_text(text):
    """把单元格里的位号串（如 'R1,R2,R3' 或 'C1,C2'）拆分为列表。"""
    items = re.split(r"[,;\s\u3001]+", str(text))
    return [x.strip().upper() for x in items if x.strip()]


# ---------- 核对逻辑 --------------------------------
class CompareResult:
    def __init__(self):
        self.rows = []          # 每行结果 dict
        self.stats = {}

    @staticmethod
    def stat(rows_designators):
        n_total = 0
        n_ok = 0
        n_bad = 0
        bad_list = []
        for row in rows_designators:
            for d in row["designators"]:
                n_total += 1
                if d["found"]:
                    n_ok += 1
                else:
                    n_bad += 1
                    bad_list.append(d)
        return n_total, n_ok, n_bad, bad_list


def check_designators(headers, data, bom_col_idx, pdf_tokens):
    """按位号核对：选中列必须是 'Part Reference' 类的位号列。"""
    col_name = headers[bom_col_idx]
    qty_idx = headers.index("Quantity") if "Quantity" in headers else None
    rows = []
    for row in data:
        rn = row["header_row"]
        raw = row["values"][bom_col_idx]
        dlist = split_designators_text(raw)
        if not dlist:
            continue
        dres = []
        for d in dlist:
            found = d in pdf_tokens
            pages = pdf_tokens.get(d, []) if found else []
            dres.append({"designator": d, "found": found, "pages": pages})
        rows.append(
            {
                "header_row": rn,
                "col_name": col_name,
                "raw": raw,
                "qty": row["values"][qty_idx] if qty_idx is not None else "",
                "designators": dres,
            }
        )
    return rows


def check_value_in_pdf(headers, data, bom_col_idx, pdf_all_words, designators_only=True):
    """全文匹配：把单元格的值作为关键词在 PDF 解码文字中查找。"""
    col_name = headers[bom_col_idx]
    rows = []
    for row in data:
        rn = row["header_row"]
        raw = row["values"][bom_col_idx]
        if raw is None or str(raw).strip() == "":
            continue
        vals = split_designators_text(raw) if designators_only else [str(raw).strip()]
        found_val = []
        for v in vals:
            hits = [p for p, y, w, x0, x1, top in pdf_all_words if w.upper() == v.upper()]
            if hits:
                found_val.append(True)
            else:
                # 子串匹配（针对值里有'/'-'/空格' 等情形）
                sub = []
                for p, y, w, x0, x1, top in pdf_all_words:
                    if v.upper() in w.upper() or w.upper() in v.upper():
                        sub.append(p)
                found_val.append(bool(sub) or bool(hits))
        rows.append(
            {
                "header_row": rn,
                "col_name": col_name,
                "raw": raw,
                "designators": [
                    {
                        "designator": v,
                        "found": any(found_val),
                        "pages": [],
                    }
                    for v in vals
                ],
            }
        )
    return rows


# ============================================================
#  对比模式：Excel vs Excel / PDF vs PDF
#  思路：把每个文件归一化成"位号集合"，再算集合差异
# ============================================================
def extract_designators_from_excel(bom_path, col_name=None):
    """从 Excel BOM 中提取全部位号(自动识别位号列)。返回 (位号集, 位号->行号)。"""
    headers, data = load_bom(bom_path)
    col = col_name or (_designator_col_name(headers) or headers[0])
    if col not in headers:
        col = headers[0]
    idx = headers.index(col)
    mapping = {}
    for row in data:
        raw = row["values"][idx]
        for d in split_designators_text(raw):
            if re.match(r"^[A-Za-z]{1,4}\d{1,4}$", d):
                mapping.setdefault(d, row["header_row"])
    return set(mapping.keys()), mapping, col


def extract_designators_from_pdf(pdf_path):
    """从 PDF 原理图中提取全部位号。返回 (位号集, 位号->页码列表)。"""
    pages_words, pages_text, _ = decode_pdf_blocks(pdf_path)
    tokens, _ = extract_designators(pages_words)
    return set(tokens.keys()), tokens


def compare_two_sets(name_a, set_a, name_b, set_b):
    """两个位号集合对比，返回 (一致, 仅A有, 仅B有) 三组。"""
    common = sorted(set_a & set_b)
    only_a = sorted(set_a - set_b)   # A有B无
    only_b = sorted(set_b - set_a)   # B有A无
    return common, only_a, only_b


# ============================================================
#  主键比对：方式A 逐位号 / 方式B 按物料分组
# ============================================================
# 单位归一化映射：0.1uF == 100nF，1K == 1000，等
_CAP_UNITS = [
    ("mf", 1000_000), ("uf", 1000_000), ("μf", 1000_000), ("µf", 1000_000),
    ("f", 1000_000_000), ("pf", 1), ("nf", 1000), ("pico", 1), ("nano", 1000),
]
_RES_UNITS = [
    ("mohm", 0.001), ("mΩ", 0.001), ("ohm", 1), ("Ω", 1), ("ω", 1),
    ("kohm", 1000), ("kω", 1000), ("k", 1000), ("mohm", 0.001),
    ("kΩ", 1000), ("kohm", 1000), ("gohm", 1000_000_000),
]


def normalize_cap_value(s):
    """把电容值归一化为 pf：0.1uF=100nF=100000pF。"""
    s = (s or "").strip().lower().replace(" ", "")
    if not s:
        return None
    s = re.sub(r"[^0-9.a-z\u00b5\u03bc]", "", s)
    if not re.search(r"\d", s):
        return s or None
    m = re.match(r"^([\d.]+)([a-z\u00b5\u03bc]*)$", s)
    if not m:
        return s
    num, unit = m.groups()  # 提取数字和单位
    # 数字在前、代号在后的解析
    try:
        val = float(num)
    except Exception:
        return s
    unit = unit or "f"
    mult = 1  # 基础单位 pF
    # 从大到小匹配（f 是最小根基，先处理 μ/m/n/p）
    table = {"pf": 1, "p": 1, "nf": 1000, "n": 1000, "uf": 1_000_000,
             "μf": 1_000_000, "µf": 1_000_000, "u": 1_000_000, "mf": 1_000_000,
             "f": 1_000_000_000}
    for u0, mm in sorted(table.items(), key=lambda kv: -len(kv[0])):
        if unit == u0 or unit.endswith(u0):
            mult = mm
            break
    return round(val * mult, 3)


def normalize_res_value(s):
    """把电阻值归一化为 ohm：1K=1000，2K4=2400，4R7=4.7，100R=100。"""
    s = (s or "").strip().lower().replace(" ", "")
    if not s:
        return None     # 4R7 / 2K4 / 1M5 简写形式
    if re.match(r"^\d+[rkm][\d.]+$", s):
        m = re.match(r"^(\d+)([rkm])([\d.]+)$", s)
        base = float(m.group(1) + "." + m.group(3))
        mult = {"r": 1, "k": 1000, "m": 1_000_000}[m.group(2)]
        return round(base * mult, 3)
    # 100R / 1K5 结尾形式
    if re.match(r"^[\d.]+[rkm]$", s):
        m = re.match(r"^([\d.]+)([rkm])$", s)
        mult = {"r": 1, "k": 1000, "m": 1_000_000}[m.group(2)]
        return round(float(m.group(1)) * mult, 3)
    if "ohm" in s or "Ω" in s or "ω" in s:
        m = re.match(r"^([\d.]+)\s*([a-z\u03a9\u03c9]*)$",
                     s.replace("mω", "m").replace("mΩ", "m").replace("kΩ", "k"))
        if not m:
            return s
        num, unit = m.groups()
        mult = 1
        table = {"pohm": 1e-6, "mohm": 1e-3, "ohm": 1, "Ω": 1, "ω": 1,
                 "kohm": 1000, "kω": 1000, "mω": 1_000_000, "mohm": 1e-3}
        for u0, mm in sorted(table.items(), key=lambda kv: -len(kv[0])):
            if unit.endswith(u0):
                mult = mm
                break
        try:
            return round(float(num) * mult, 3)
        except Exception:
            return s
    # "10k" "100k" "1k" 等
    m = re.match(r"^([\d.]+)\s*([km]?)$", s)
    if m:
        num, unit = m.groups()
        try:
            mult = {"": 1, "k": 1000, "m": 1_000_000}.get(unit, 1)
            return round(float(num) * mult, 3)
        except Exception:
            return s
    return s


def smart_norm_value(v):
    """通用值归一化：数字+单位换算为统一结果；否则返回清理后字符串。"""
    s = str(v or "").strip()
    if not s:
        return None
    s2 = s.lower()

    # 电容类：F / uF / nF / pF
    if any(unit in s2 for unit in ["uf", "μf", "µf", "nf", "pf", "mf"]):
        r = normalize_cap_value(s)
        if r is not None:
            return r
    # 电阻类：Ω/ohm/K(欧)/R 缩写
    if (any(unit in s2 for unit in ["ohm", "Ω", "ω", "k", "r"]) and
            "f" not in s2 and "uf" not in s2 and "µ" not in s2 and "μ" not in s2):
        r = normalize_res_value(s)
        if r is not None and r != s:
            return r
    # 纯数字：统一成数值（1 vs 1.0）
    if re.fullmatch(r"[0-9.]+", s):
        try:
            return float(s)
        except Exception:
            return s
    return s


def _px(a, b):
    """比较两个字段值（支持归一化），返回 (是否一致, 规范值a, 规范值b)。"""
    na, nb = smart_norm_value(a), smart_norm_value(b)
    if na is None and nb is None:
        return True, "", ""
    eq = (na == nb) or (str(na).upper() == str(nb).upper())
    return eq, "" if na is None else str(na), "" if nb is None else str(nb)


def _footprint_family(d):
    """把封装名归并成大类，方便比较不同写法。"""
    s = str(d or "").strip().lower().replace(" ", "")
    fam = {
        "capc1005": "0402", "capc0603": "0201", "capc1608": "0603", "capc2012": "0805",
        "capc3216": "1206", "resc1005": "0402", "resc0603": "0201", "resc1608": "0603",
        "resc2013": "0805", "indc1005": "0402", "res": "",
        "0402": "0402", "0201": "0201", "0603": "0603", "0805": "0805", "1206": "1206",
    }
    for prefix, f in fam.items():
        if s.startswith(prefix):
            return f
    return s or ""


def load_excel_table(path_a, path_b):
    """加载两个 Excel，返回 (表头A, 行A, 表头B, 行B)。"""
    ha, da = load_bom(path_a)
    hb, db = load_bom(path_b)
    return ha, da, hb, db


def compare_excel_excel_mpn(path_a, path_b, key_a=None, key_b=None, qty_a=None, qty_b=None):
    """Excel vs Excel 以 用户指定的一列或多列主键 比对：
      ① 双方各自可用「多个主键列」（如 MPN+料号、或 MPN+零件号），组合成多列主键匹配；
         key_a/key_b 可为字符串(单列)或 列表(多列)，空/None 时自动识别 MPN(物料号) 列；
      ② 主键一致后再比数量(qty_a/qty_b)/位号/厂商等字段。
      返回 rows(含 'item','status','field_diffs') 与 stats(mode='excelmpn')。
    """
    ha, da, hb, db = load_excel_table(path_a, path_b)

    def _resolve_multi(headers, cols, keys):
        """把用户选的列(可为单列字符串或多列列表)解析成 headers 中存在的列名列表。"""
        if cols is None:
            cols = []
        if isinstance(cols, str):
            cols = [cols]
        cols = [c for c in cols if c and c in headers]
        if cols:
            return cols
        auto = _find_col(headers, keys)
        return [auto] if auto else ([headers[0]] if headers else [])

    ma = _resolve_multi(ha, key_a, _MPN_COL_KEYS)      # 文件A 主键列(可多个)
    mb = _resolve_multi(hb, key_b, _MPN_COL_KEYS)      # 文件B 主键列(可多个)
    qa = _find_col(ha, _QTY_COL_KEYS)                  # 文件A 数量列(自动)
    qb = _find_col(hb, _QTY_COL_KEYS)                  # 文件B 数量列(自动)
    if qty_a and qty_a in ha:
        qa = qty_a
    if qty_b and qty_b in hb:
        qb = qty_b
    ra = _find_col(ha, _REF_COL_KEYS)                  # 文件A 位号列
    rb = _find_col(hb, _REF_COL_KEYS)                  # 文件B 位号列
    fa = _find_col(ha, _MFG_COL_KEYS)                  # 文件A 厂商列
    fb = _find_col(hb, _MFG_COL_KEYS)                  # 文件B 厂商列
    ia_ = _find_col(ha, _ITM_COL_KEYS)                 # 文件A 料号页号列(ITEM NUMBER)
    ib_ = _find_col(hb, _ITM_COL_KEYS)                 # 文件B 料号页号列
    da_ = _find_col(ha, _DESC_COL_KEYS)                # 文件A 描述列(ITEM NAME)
    db_ = _find_col(hb, _DESC_COL_KEYS)                # 文件B 描述列
    if not ma or not mb:
        raise RuntimeError("文件A/文件B 中找不到主键列（%s / %s）" % (ma, mb))

    def build(headers, data, mcol_list, qcol, refcol, mfgcol, itmcol, desccol):
        mi = [headers.index(m) for m in mcol_list]
        qi = headers.index(qcol) if qcol in headers else None
        ri = headers.index(refcol) if refcol in headers else None
        fi = headers.index(mfgcol) if mfgcol in headers else None
        ii = headers.index(itmcol) if itmcol in headers else None
        di = headers.index(desccol) if desccol in headers else None
        idx = {}
        for row in data:
            pk = tuple(str(row["values"][i] or "").strip().upper() for i in mi)
            if not any(pk):
                continue
            idx.setdefault(pk, []).append({
                "row": row, "qty": (row["values"][qi] if qi is not None else "") or "",
                "refs": split_designators_text(row["values"][ri]) if ri is not None else [],
                "mfg": (row["values"][fi] if fi is not None else "") or "",
                "itm": (row["values"][ii] if ii is not None else "") or "",
                "desc": (row["values"][di] if di is not None else "") or "",
            })
        return idx

    def _key_label(pk, mcol_list):
        """把组合主键显示成可读字符串（如 'MPN=T598B... / IPN=E-120-000132'）。"""
        if len(mcol_list) == 1:
            return str(pk[0])
        return " | ".join("%s=%s" % (m, str(v)) for m, v in zip(mcol_list, pk) if v)

    ia = build(ha, da, ma, qa, ra, fa, ia_, da_)
    ib = build(hb, db, mb, qb, rb, fb, ib_, db_)
    keys = sorted(set(ia.keys()) | set(ib.keys()), key=lambda k: tuple(k))

    rows = []
    n_ok = n_diff = n_only_a = n_only_b = 0
    for k in keys:
        label = _key_label(k, ma if k in ia else mb)
        ra_ = ia.get(k, [])
        rb_ = ib.get(k, [])
        if not rb_:
            n_only_a += 1
            rows.append({"item": label, "key": k, "status": "仅文件A", "rec_a": ra_[0], "rec_b": None,
                         "field_diffs": [], "hdra": ha, "hdrb": hb})
            continue
        if not ra_:
            n_only_b += 1
            rows.append({"item": label, "key": k, "status": "仅文件B", "rec_a": None, "rec_b": rb_[0],
                         "field_diffs": [], "hdra": ha, "hdrb": hb})
            continue
        # 主键一致：逐个比较 数量 / 位号 / 厂商（取两侧首个记录做代表）
        ra0, rb0 = ra_[0], rb_[0]
        diffs = []
        qa_, qb_ = str(ra0["qty"] or "").strip(), str(rb0["qty"] or "").strip()
        if qa_ != qb_ and qa_ and qb_:
            diffs.append({"field": "数量(%s)" % (qa or "Quantity"), "a": qa_, "b": qb_, "raw_a": qa_, "raw_b": qb_})
        sa, sb = set(ra0["refs"]), set(rb0["refs"])
        if sa != sb:
            diffs.append({"field": "位号(%s)" % (ra or "Reference"),
                          "a": ",".join(sorted(sa)), "b": ",".join(sorted(sb)),
                          "raw_a": ",".join(sorted(sa)), "raw_b": ",".join(sorted(sb))})
        mfa_, mfb_ = str(ra0["mfg"] or "").strip(), str(rb0["mfg"] or "").strip()
        mfa_n = re.sub(r"[\s\-_./]", "", mfa_.upper())
        mfb_n = re.sub(r"[\s\-_./]", "", mfb_.upper())
        if mfa_n != mfb_n and mfa_ and mfb_:
            diffs.append({"field": "厂商(Manufacturer)", "a": mfa_, "b": mfb_,
                          "raw_a": mfa_, "raw_b": mfb_})
        if diffs:
            n_diff += 1
            rows.append({"item": label, "key": k, "status": "字段不一致", "rec_a": ra0, "rec_b": rb0,
                         "field_diffs": diffs, "hdra": ha, "hdrb": hb})
        else:
            n_ok += 1
            rows.append({"item": label, "key": k, "status": "一致", "rec_a": ra0, "rec_b": rb0,
                         "field_diffs": [], "hdra": ha, "hdrb": hb})

    stats = {
        "all": {
            "文件A主键列": ", ".join(ma),
            "文件B主键列": ", ".join(mb),
            "文件A数量列": qa or "(未识别)",
            "文件B数量列": qb or "(未识别)",
            "MPN/主键总数": len(keys),
            "一致(数量/位号/厂商匹配)": n_ok,
            "字段不一致": n_diff,
            "仅文件A有": n_only_a,
            "仅文件B有": n_only_b,
        },
        "file_a": path_a,
        "file_b": path_b,
        "bad": [{"designator": r["item"], "row": ""} for r in rows if r["status"] != "一致"],
        "extra": [r["item"] for r in rows if r["status"] == "仅文件B"],
        "in_bom": {r["item"] for r in rows if r["status"] != "仅文件B"},
        "common": [r["item"] for r in rows if r["status"] == "一致"],
        "only_a": [r["item"] for r in rows if r["status"] == "仅文件A"],
        "only_b": [r["item"] for r in rows if r["status"] == "仅文件B"],
        "mode": "excelmpn",
    }
    return rows, stats


def compare_excel_group(path_a, path_b, group_fields=("Value", "Manufacturer PN"),
                        qty_col="Quantity", ref_col="Part Reference"):
    """方式B：按物料分组比对。
    以 group_fields 为键分组，比较每组位号集合与数量。
    """
    ha, da, hb, db = load_excel_table(path_a, path_b)

    def group_index(headers, data):
        g = {}
        for row in data:
            vals = row["values"]
            key = tuple((smart_norm_value(vals[headers.index(f)]) if f in headers
                         else None) for f in group_fields)
            # 料号空值：把 None 视为能匹配
            key = tuple("" if k is None else str(k) for k in key)
            refs = split_designators_text(vals[headers.index(ref_col)])
            g.setdefault(key, []).extend(refs)
        return g

    ga = group_index(ha, da)
    gb = group_index(hb, db)
    all_keys = sorted(set(ga.keys()) | set(gb.keys()))

    rows = []
    n_ok, n_diff, n_only_a, n_only_b = 0, 0, 0, 0
    for key in all_keys:
        ra = sorted(ga.get(key, []))
        rb = sorted(gb.get(key, []))
        if not rb:
            n_only_a += 1
            rows.append({"item": " | ".join(str(k) for k in key), "status": "仅文件A",
                         "refs_a": ra, "refs_b": [], "qty_a": len(ra), "qty_b": 0})
            continue
        if not ra:
            n_only_b += 1
            rows.append({"item": " | ".join(str(k) for k in key), "status": "仅文件B",
                         "refs_a": [], "refs_b": rb, "qty_a": 0, "qty_b": len(rb)})
            continue
        # 同组：比较位号集合
        only_in_a = sorted(set(ra) - set(rb))
        only_in_b = sorted(set(rb) - set(ra))
        if only_in_a or only_in_b:
            n_diff += 1
            rows.append({"item": " | ".join(str(k) for k in key), "status": "位号集合差异",
                         "refs_a": ra, "refs_b": rb, "qty_a": len(ra), "qty_b": len(rb),
                         "only_a": only_in_a, "only_b": only_in_b})
        else:
            n_ok += 1
            rows.append({"item": " | ".join(str(k) for k in key), "status": "一致",
                         "refs_a": ra, "refs_b": rb, "qty_a": len(ra), "qty_b": len(rb)})

    stats = {
        "all": {
            "分组键": "、".join(group_fields),
            "物料分组数": len(all_keys),
            "一致": n_ok,
            "位号集合差异": n_diff,
            "仅文件A有": n_only_a,
            "仅文件B有": n_only_b,
        },
        "bad": [{"designator": r["item"], "row": ""} for r in rows
                if r["status"] in ("位号集合差异", "仅文件A")],
        "extra": [r["item"] for r in rows if r["status"] == "仅文件B"],
        "in_bom": {r["item"] for r in rows if r["status"] != "仅文件B"},
        "common": [r["item"] for r in rows if r["status"] == "一致"],
        "only_a": [r["item"] for r in rows if r["status"] == "仅文件A"],
        "only_b": [r["item"] for r in rows if r["status"] == "仅文件B"],
        "mode": "group",
    }
    return rows, stats


def compare_excel_excel(path_a, path_b, col_name):
    """两个 Excel BOM 位号对比。"""
    set_a, map_a, col_a = extract_designators_from_excel(path_a, col_name)
    set_b, map_b, col_b = extract_designators_from_excel(path_b, col_name)
    common, only_a, only_b = compare_two_sets(os.path.basename(path_a), set_a,
                                              os.path.basename(path_b), set_b)
    rows = []
    for d in common:
        rows.append({"item": d, "status": "一致", "a": map_a.get(d, ""), "b": map_b.get(d, ""),
                     "a_pages": "", "b_pages": ""})
    for d in only_a:
        rows.append({"item": d, "status": "仅文件A", "a": map_a.get(d, ""), "b": "",
                     "a_pages": "", "b_pages": ""})
    for d in only_b:
        rows.append({"item": d, "status": "仅文件B", "a": "", "b": map_b.get(d, ""),
                     "a_pages": "", "b_pages": ""})
    stats = {
        "all": {
            "文件A位号数": len(set_a),
            "文件B位号数": len(set_b),
            "两文件一致": len(common),
            "仅文件A有": len(only_a),
            "仅文件B有": len(only_b),
            "对比列": col_a,
        },
        "bad": [{"designator": d, "row": map_a.get(d, "")} for d in only_a],
        "extra": only_b,
        "in_bom": set_a,
        "common": common,
        "only_a": only_a,
        "only_b": only_b,
    }
    return rows, stats


def compare_pdf_pdf(path_a, path_b):
    """两个 PDF 原理图位号对比。"""
    set_a, tokens_a = extract_designators_from_pdf(path_a)
    set_b, tokens_b = extract_designators_from_pdf(path_b)
    common, only_a, only_b = compare_two_sets(os.path.basename(path_a), set_a,
                                              os.path.basename(path_b), set_b)
    rows = []
    for d in common:
        rows.append({"item": d, "status": "一致",
                     "a": "", "b": "",
                     "a_pages": ",".join(str(p) for p in tokens_a.get(d, [])),
                     "b_pages": ",".join(str(p) for p in tokens_b.get(d, []))})
    for d in only_a:
        rows.append({"item": d, "status": "仅文件A", "a": "", "b": "",
                     "a_pages": ",".join(str(p) for p in tokens_a.get(d, [])), "b_pages": ""})
    for d in only_b:
        rows.append({"item": d, "status": "仅文件B", "a": "", "b": "",
                     "a_pages": "", "b_pages": ",".join(str(p) for p in tokens_b.get(d, []))})
    stats = {
        "all": {
            "PDF A 位号数": len(set_a),
            "PDF B 位号数": len(set_b),
            "两文件一致": len(common),
            "仅 PDF A 有": len(only_a),
            "仅 PDF B 有": len(only_b),
        },
        "bad": [{"designator": d, "row": ""} for d in only_a],
        "extra": only_b,
        "in_bom": set_a,
        "common": common,
        "only_a": only_a,
        "only_b": only_b,
    }
    return rows, stats


# ---------- 报告生成 --------------------------------
def gen_report(bom_path, rows, stats, pdf_tokens, pages_text, out_dir, formats=None):
    """生成报告。formats 为导出类型集合，可选 {"xlsx","docx","pdf","txt"}；None=全部。"""
    if formats is None:
        formats = {"xlsx", "docx", "pdf", "txt"}
    else:
        formats = set(formats)
    ts = time.strftime("%Y%m%d_%H%M%S")
    # 根据 row 结构自动识别模式：对比模式 rows 含 "item"/"status" 且无 "designators"
    is_compare = bool(rows) and "designators" not in rows[0]
    prefix = "对比报告" if is_compare else "BOM核对报告"
    xlsx_path = os.path.join(out_dir, f"{prefix}_{ts}.xlsx")
    txt_path = os.path.join(out_dir, f"{prefix}_{ts}.txt")
    docx_path = os.path.join(out_dir, f"{prefix}_{ts}.docx")

    wb = openpyxl.Workbook()
    ws = wb.active
    if is_compare:
        mode = stats.get("mode", "set")
        if mode == "mpn":
            ws.title = "MPN 对比(料号)"
            ws.append(["MPN(料号)", "核对结果", "Excel数量", "PDF数量", "Excel位号",
                       "PDF位号", "厂商", "PDF页"])
            for row in rows:
                ws.append([row["item"], row["status"], row["qty_a"], row["qty_b"],
                           row["refs_a"], row["refs_b"], row["mfg"], row["pages"]])
        elif mode == "excelmpn":
            from openpyxl.styles import PatternFill, Font
            ws.title = "Excel MPN 对比"
            red_fill = PatternFill("solid", fgColor="FFC7CE")
            red_font = Font(color="9C0006", bold=True)
            ws.append(["MPN(物料号)", "结果", "数量A", "数量B", "位号A", "位号B",
                       "厂商A", "厂商B", "差异明细"])
            for row in rows:
                ra = row.get("rec_a") or {}
                rb = row.get("rec_b") or {}
                desc = "；".join("%s: %s → %s" % (d["field"], d["a"] or "空", d["b"] or "空")
                                 for d in row["field_diffs"])
                cell = ws.append([row["item"], row["status"],
                                  ra.get("qty", ""), rb.get("qty", ""),
                                  ",".join(ra.get("refs", [])), ",".join(rb.get("refs", [])),
                                  ra.get("mfg", ""), rb.get("mfg", ""), desc])
                # 器件已使用/未使用、差异项 → 红色标记
                if row["status"] != "一致":
                    r_ = ws.max_row
                    for cc in range(1, ws.max_column + 1):
                        c_ = ws.cell(row=r_, column=cc)
                        c_.fill = red_fill
                        c_.font = red_font
        elif mode == "pdf2excel":
            ws.title = "PDF→Excel 器件核对"
            ws.append(["PDF位号", "核对结果", "Excel值(Value)", "Excel封装(Footprint)",
                       "Excel数量", "PDF附近标注"])
            for row in rows:
                ws.append([row["item"], row["status"], row["valueA"], row["footA"],
                           row["qty"], row["near"]])
        else:
            ws.title = "对比明细"
            ws.append(["器件/位号", "结果", "文件A行号/页码", "文件B行号/页码"])
            for row in rows:
                ws.append([row["item"], row["status"], row["a_pages"] or row["a"],
                           row["b_pages"] or row["b"]])
    else:
        ws.title = "核对明细"
        ws.append(
            ["BOM行号", "列", "原值", "位号/值", "PDF中是否存在", "所在页码", "结果"]
        )
        for row in rows:
            for d in row["designators"]:
                ws.append(
                    [
                        row["header_row"],
                        row["col_name"],
                        row["raw"],
                        d["designator"],
                        "是" if d["found"] else "否",
                        ",".join(str(p) for p in d["pages"]) if d["found"] else "",
                        "一致" if d["found"] else "不一致",
                    ]
                )

    # Sheet2: 汇总
    ws2 = wb.create_sheet("汇总")
    ws2.append(["统计项", "数量"])
    for k, v in stats["all"].items():
        ws2.append([k, v])
    ws2.append([])
    if is_compare:
        ws2.append(["仅文件A有：", len(stats["bad"])])
        for b in stats["bad"]:
            ws2.append(["", b["designator"]])
        ws2.append(["仅文件B有：", len(stats["extra"])])
        for e in stats["extra"]:
            ws2.append(["", e])
    else:
        ws2.append(["BOM未在PDF中找到的位号：", len(stats["bad"])])
        for b in stats["bad"]:
            ws2.append(["", b["designator"], "", "不一致"])
        # Sheet3: PDF 中出现的所有器件位号
        ws3 = wb.create_sheet("PDF全部位号")
        ws3.append(["位号", "出现页码", "是否在BOM"])
        in_bom = stats.get("in_bom", set())
        for t in sorted(pdf_tokens.keys(),
                        key=lambda s: (s[0], int(re.search(r"\d+", s).group()) if re.search(r"\d+", s) else 0)):
            ws3.append([t, ",".join(str(p) for p in pdf_tokens[t]), "是" if t in in_bom else "否"])
    if "xlsx" in formats:
        wb.save(xlsx_path)
    else:
        xlsx_path = None

    # txt 报告
    if "txt" in formats:
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write("=" * 70 + "\n")
            if is_compare:
                f.write("文件对比报告（Excel-Excel / Excel-PDF / PDF-PDF）\n")
            else:
                f.write("BOM 与 PDF 原理图 器件核对报告\n")
            f.write("来源: %s\n" % bom_path)
            f.write("时间: %s\n" % time.strftime("%Y-%m-%d %H:%M:%S"))
            f.write("=" * 70 + "\n\n")
            f.write("【汇总】\n")
            for k, v in stats["all"].items():
                f.write("  %-30s: %s\n" % (k, v))
            if is_compare:
                mode = stats.get("mode", "set")
                if mode == "mpn":
                    f.write("\n【MPN 对比明细】\n")
                    for row in rows:
                        f.write("  [%s] %s  厂商=%s  数量 Excel=%s / PDF=%s\n"
                                % (row["status"], row["item"], row["mfg"], row["qty_a"], row["qty_b"]))
                        f.write("        Excel位号: %s\n" % row["refs_a"])
                        f.write("        PDF位号 : %s\n" % row["refs_b"])
                        for mm in row.get("mismatchs", []):
                            f.write("        差异: %s\n" % mm)
                    f.write("\n【PDF有而Excel无】\n")
                    for e in sorted(stats["extra"]):
                        f.write("  %s\n" % e)
                elif mode == "excelmpn":
                    f.write("\n【Excel MPN 对比明细】\n")
                    for row in rows:
                        ra = row.get("rec_a") or {}
                        rb = row.get("rec_b") or {}
                        f.write("  [%s] %s\n" % (row["status"], row["item"]))
                        f.write("        数量: A=%s  B=%s\n" % (ra.get("qty", ""), rb.get("qty", "")))
                        f.write("        位号: A=%s\n                B=%s\n"
                                % (",".join(ra.get("refs", [])), ",".join(rb.get("refs", []))))
                        f.write("        厂商: A=%s  B=%s\n" % (ra.get("mfg", ""), rb.get("mfg", "")))
                        for d in row["field_diffs"]:
                            f.write("        差异[%s]: %s → %s\n" % (d["field"], d["a"] or "(空)", d["b"] or "(空)"))
                    f.write("\n【仅文件A有】\n")
                    for b in stats["bad"]:
                        f.write("  %s\n" % b["designator"])
                    f.write("\n【仅文件B有】\n")
                    for e in sorted(stats["extra"]):
                        f.write("  %s\n" % e)
                elif mode == "pdf2excel":
                    f.write("\n【疑似不一致 / 待确认】\n")
                    for b in stats["bad"]:
                        f.write("  %s\n" % b["designator"])
                    f.write("\n【PDF有Excel无】\n")
                    for t in sorted(stats["extra"]):
                        f.write("  %s\n" % t)
                    f.write("\n【明细】\n")
                    for row in rows:
                        f.write("  [%s] %s  Excel值=%s 封装=%s  附近标注=%s\n"
                                % (row["status"], row["item"], row["valueA"], row["footA"], row["near"]))
                else:
                    f.write("\n【仅文件A有】\n")
                    for b in stats["bad"]:
                        f.write("  %s\n" % b["designator"])
                    f.write("\n【仅文件B有】\n")
                    for t in sorted(stats["extra"]):
                        f.write("  %s\n" % t)
                    f.write("\n【明细】\n")
                    for row in rows:
                        f.write("  %-6s | %-10s | A:%s | B:%s\n"
                                % (row["status"], row["item"], row["a_pages"] or row["a"],
                                   row["b_pages"] or row["b"]))
            else:
                f.write("\n【不一致 - BOM位号在PDF中未找到】\n")
                for b in stats["bad"]:
                    f.write("  %s  (BOM第%d行)\n" % (b["designator"], b["row"]))
                f.write("\n【PDF中出现但不在BOM的位号】\n")
                for t in sorted(stats["extra"]):
                    f.write("  %s  页码:%s\n" % (t, ",".join(str(p) for p in pdf_tokens[t])))
                f.write("\n【明细】\n")
                for row in rows:
                    for d in row["designators"]:
                        mark = "OK " if d["found"] else "MISS"
                        f.write(
                            "%s | 行%-3d | %-6s | %-10s | %s\n"
                            % (mark, row["header_row"], row["col_name"], d["designator"], row["raw"])
                        )
    else:
        txt_path = None

    if "txt" in formats:
        _export_html_report(txt_path, stats, rows, is_compare, pdf_tokens)
    # 生成 Arena 风格 PDF 对比报告（仅 Excel-Excel / MPN 模式）
    if "pdf" in formats:
        base = os.path.splitext(os.path.basename(
            txt_path or xlsx_path or docx_path or "rep_%s" % ts))[0]
        _export_pdf_report(out_dir, base, rows, stats,
                           stats.get("file_a") or bom_path, stats.get("file_b") or "")
    if "docx" in formats:
        _export_word_report(docx_path, rows, stats, is_compare)
    return xlsx_path, txt_path, docx_path


def _esc_html(s):
    return (str(s or "").replace("&", "&amp;").replace("<", "&lt;")
            .replace(">", "&gt;").replace('"', "&quot;"))


def _export_html_report(txt_path, stats, rows, is_compare, pdf_tokens):
    """额外生成同名的 .html 报告，方便浏览器查看/归档。"""
    html_path = os.path.splitext(txt_path)[0] + ".html"
    try:
        with open(html_path, "w", encoding="utf-8") as f:
            f.write("<!DOCTYPE html><html lang='zh'><head><meta charset='utf-8'>")
            f.write("<title>BOM 核对报告</title>")
            f.write("<style>body{font-family:'Microsoft YaHei',sans-serif;margin:24px;color:#23303B}"
                    "h1{color:#1E4E84}table{border-collapse:collapse;width:100%}"
                    "td,th{border:1px solid #dfe6f0;padding:6px 10px;font-size:13px;text-align:left}"
                    "th{background:#E8F0FA;color:#1E4E84}.ok{background:#E7F7EA;color:#1E7B34}"
                    ".bad{background:#FDEAEA;color:#C0392B}.warn{background:#FFF4E0;color:#996A00}</style>")
            f.write("</head><body>")
            f.write("<h1>BOM 核对 / 对比报告</h1>")
            f.write("<p>生成时间: %s</p>" % time.strftime("%Y-%m-%d %H:%M:%S"))
            f.write("<h2>汇总</h2><table><tr><th>统计项</th><th>数量</th></tr>")
            for k, v in stats.get("all", {}).items():
                f.write("<tr><td>%s</td><td>%s</td></tr>" % (_esc_html(k), _esc_html(v)))
            f.write("</table>")
            if is_compare:
                if stats.get("mode") == "mpn":
                    f.write("<h2>MPN 对比明细</h2><table><tr><th>MPN</th><th>结果</th>"
                            "<th>Excel数量</th><th>PDF数量</th><th>Excel位号</th><th>PDF位号</th><th>厂商</th></tr>")
                    for row in rows:
                        cls = "ok" if row.get("status") == "一致" else "bad"
                        f.write("<tr class='%s'><td>%s</td><td>%s</td><td>%s</td><td>%s</td>"
                                "<td>%s</td><td>%s</td><td>%s</td></tr>"
                                % (cls, _esc_html(row.get("item")), _esc_html(row.get("status")),
                                   _esc_html(row.get("qty_a")), _esc_html(row.get("qty_b")),
                                   _esc_html(row.get("refs_a")), _esc_html(row.get("refs_b")),
                                   _esc_html(row.get("mfg"))))
                    f.write("</table>")
                    f.write("<h2>PDF有而Excel无</h2><table><tr><th>料号</th></tr>")
                    for e in stats.get("extra", []):
                        f.write("<tr class='warn'><td>%s</td></tr>" % _esc_html(e))
                    f.write("</table>")
                elif stats.get("mode") == "excelmpn":
                    # Excel vs Excel MPN 对比（红色高亮已使用/未使用/差异器件）
                    f.write("<h2>Excel MPN 对比明细</h2>")
                    f.write("<table><tr><th>MPN(物料号)</th><th>结果</th><th>数量A</th><th>数量B</th>"
                            "<th>位号A</th><th>位号B</th><th>厂商A</th><th>厂商B</th><th>差异</th></tr>")
                    for row in rows:
                        ra = row.get("rec_a") or {}
                        rb = row.get("rec_b") or {}
                        cls = "ok" if row.get("status") == "一致" else "bad"
                        desc = "；".join("%s: %s→%s" % (d["field"], d["a"] or "空", d["b"] or "空")
                                         for d in row["field_diffs"])
                        f.write("<tr class='%s'><td>%s</td><td>%s</td><td>%s</td><td>%s</td>"
                                "<td>%s</td><td>%s</td><td>%s</td><td>%s</td><td>%s</td></tr>"
                                % (cls, _esc_html(row.get("item")), _esc_html(row.get("status")),
                                   _esc_html(ra.get("qty")), _esc_html(rb.get("qty")),
                                   _esc_html(",".join(ra.get("refs", []))),
                                   _esc_html(",".join(rb.get("refs", []))),
                                   _esc_html(ra.get("mfg")), _esc_html(rb.get("mfg")),
                                   _esc_html(desc)))
                    f.write("</table>")
                    f.write("<h2>仅文件A有 / 仅文件B有（红色=已使用/未使用）</h2>")
                    f.write("<table><tr><th>类别</th><th>MPN</th></tr>")
                    for b in stats.get("bad", []):
                        f.write("<tr class='bad'><td>仅文件A有</td><td>%s</td></tr>" % _esc_html(b["designator"]))
                    for e in stats.get("extra", []):
                        f.write("<tr class='bad'><td>仅文件B有</td><td>%s</td></tr>" % _esc_html(e))
                    f.write("</table>")
                else:
                    f.write("<h2>不一致 / 待确认</h2><table><tr><th>位号</th><th></th></tr>")
                    for b in stats.get("bad", []):
                        f.write("<tr class='bad'><td>%s</td><td>%s</td></tr>"
                                % (_esc_html(b["designator"]), _esc_html(b.get("row") or "")))
                    f.write("</table>")
                    f.write("<h2>仅文件B有</h2><table><tr><th>位号</th></tr>")
                    for e in stats.get("extra", []):
                        f.write("<tr class='warn'><td>%s</td></tr>" % _esc_html(e))
                    f.write("</table>")
                    f.write("<h2>明细</h2><table><tr><th>结果</th><th>器件/位号</th><th>详情</th></tr>")
                    for row in rows:
                        des = row.get("item", "")
                        st = row.get("status", "")
                        cls = "ok" if st == "一致" else "bad"
                        detail = _esc_html(row.get("valueA") or "") + " " + _esc_html(row.get("footA") or "")
                        if not detail:
                            try:
                                detail = _esc_html(",".join(row.get("refs_a", []) or row.get("a", ""))) + \
                                         " → " + _esc_html(",".join(row.get("refs_b", []) or row.get("b", "")))
                            except Exception:
                                detail = ""
                        f.write("<tr class='%s'><td>%s</td><td>%s</td><td>%s</td></tr>"
                                % (cls, _esc_html(st), _esc_html(des), detail))
                    f.write("</table>")
            else:
                f.write("<h2>明细</h2><table><tr><th>BOM行</th><th>列</th>"
                        "<th>位号/值</th><th>结果</th><th>页码</th><th>原值</th></tr>")
                for row in rows:
                    for d in row.get("designators", []):
                        cls = "ok" if d.get("found") else "bad"
                        f.write("<tr class='%s'><td>%s</td><td>%s</td><td>%s</td><td>%s</td><td>%s</td><td>%s</td></tr>"
                                % (_esc_html(row.get("header_row")), _esc_html(row.get("col_name")),
                                   _esc_html(d.get("designator")),
                                   "一致" if d.get("found") else "不一致",
                                   _esc_html(",".join(str(p) for p in d.get("pages", [])) if d.get("found") else ""),
                                   _esc_html(row.get("raw"))))
                f.write("</table>")
            f.write("</body></html>")
    except Exception:
        pass


# ---------- PDF 对比报告（仿 Arena "BOM Compare" 样式） ----------------
def _export_pdf_report(out_dir, base_name, rows, stats, path_a, path_b):
    """生成与 835&685对比.pdf 同风格的 PDF 对比报告。
    布局：A3 横向、表头 # / ITEM NUMBER / ITEM NAME / PCS / REF DES / SUBSTITUTES，
    已使用/未使用(差异)器件用红色标记(0.6,0,0)加 '>>' 前缀。"""
    pdf_path = os.path.join(out_dir, base_name + ".pdf")
    try:
        import pymupdf
    except ImportError:
        return None
    mode = stats.get("mode", "set") if stats else "set"
    if mode not in ("excelmpn", "mpn"):
        return None

    # 页面尺寸 A3 横向（与参考文档一致的点数）
    page_w, page_h = 1191.0, 842.0
    RED = (0.6, 0, 0)           # 0x990000
    BLACK = (0, 0, 0)
    BODY = "helv"               # 笔画字体（参考文档为 SegoeUI，这里用内建 sans-serif 近似）
    BOLD = "hebo"
    ITAL = "hebi"

    doc = pymupdf.open()
    page = doc.new_page(width=page_w, height=page_h)

    def _p(x, y, text, size=8, font=BODY, color=BLACK):
        page.insert_text((x, y), text, fontsize=size, fontname=font, color=color)

    # ---- 顶部页眉（仿参考文档） ----
    _p(34, 55, "Seyond Inc.", 11.2, BOLD)
    _p(34, 68, "Seyond Production", 9, BOLD)
    _p(34, 80, "Printed by", 6, BOLD)
    _p(34, 90, "Printed on %s" % time.strftime("%m/%d/%Y"), 6, BOLD)
    _p(34, 100, "Local time zone (GMT+08:00) China Standard Time", 6, BOLD)
    _p(34, 112, "Page  >  Bill of Materials  >  Compare", 8)
    _p(34, 121, "All information contained in this document is proprietary and confidential.",
       6, ITAL, (0, 0.067, 0.09))

    # ---- 文件标题（两个 BOM 版本） ----
    name_a = os.path.basename(path_a or "")
    name_b = os.path.basename(path_b or "")
    _p(50, 145, name_a, 13.5, BOLD)
    _p(50, 163, "Rev A1", 8)
    _p(680, 145, name_b, 13.5, BOLD)
    _p(680, 163, "Rev A1", 8)

    # ---- 表头 ----
    y0 = 251.0
    _p(35, y0, ">>", 8, BODY, RED)
    _p(48, y0, "#", 8, BOLD)
    _p(115, y0, "ITEM NUMBER", 8, BOLD)
    _p(284, y0, "ITEM NAME", 8, BOLD)
    _p(647, y0 - 4, "PCS", 7.5, BOLD)
    _p(743, y0, "REF DES", 8, BOLD)
    _p(1103, y0, "SUBSTITUTES", 8, BOLD)

    page.draw_line(pymupdf.Point(30, y0 + 6), pymupdf.Point(1165, y0 + 6), color=(0.8, 0.8, 0.8), width=0.6)

    # ---- 数据行 ----
    row_h = 15.0
    y = y0 + 16
    n = 0
    for row in rows:
        if y > page_h - 40:
            page = doc.new_page(width=page_w, height=page_h)
            y = 40
        n += 1
        status = row.get("status", "")
        changed = status != "一致"
        ra = row.get("rec_a") or {}
        rb = row.get("rec_b") or {}
        c = RED if changed else BLACK
        if changed:
            _p(35, y, ">>", 8, BOLD, RED)
        _p(48, y, "%02d" % n, 8, BODY, c)
        # ITEM NUMBER（显示料号/IPN，优先 item 字段）
        itm_a = str(ra.get("itm") or "") or str(row.get("item"))
        _p(115, y, _clip(itm_a, 34), 8, BODY, c)
        # ITEM NAME（描述）
        _p(284, y, _clip(str(ra.get("desc") or ""), 70), 8, BODY, c)
        # PCS 数量
        qa = str(ra.get("qty") or "")
        _p(647, y, ("PCS%s" % qa) if qa else "PCS", 8, BODY, c)
        # REF DES 位号
        refs = ra.get("refs") or []
        _p(743, y, _clip(",".join(refs), 55), 8, BODY, c)
        # SUBSTITUTES（差异简述）
        if status == "仅文件A":
            sub = "仅文件A(文件B无此料号)"
        elif status == "仅文件B":
            sub = "仅文件B(文件A无此料号)"
        elif row.get("field_diffs"):
            sub = "；".join("%s: %s→%s" % (d["field"], d["a"] or "空", d["b"] or "空")
                            for d in row["field_diffs"])[:60]
        else:
            sub = ""
        _p(1103, y, _clip(sub, 45), 8, BODY, RED if sub else BLACK)
        y += row_h

    doc.save(pdf_path)
    doc.close()
    return pdf_path


def _clip(s, n):
    s = str(s or "")
    return s if len(s) <= n else s[:n] + "…"


def _export_word_report(docx_path, rows, stats, is_compare):
    """生成 Word(.docx) 报告，已使用/未使用/差异器件红色标记。"""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.oxml.ns import qn
    except ImportError:
        return None

    def _set_font(run, size=10, bold=False, color=None):
        run.font.size = Pt(size)
        run.font.bold = bold
        if color is not None:
            run.font.color.rgb = RGBColor(*color)
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    doc = Document()
    doc.add_heading("BOM 核对 / 对比报告", level=0)
    p = doc.add_paragraph()
    _set_font(p.add_run("生成时间: %s" % time.strftime("%Y-%m-%d %H:%M:%S")), size=9)

    doc.add_heading("汇总", level=1)
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    for c, txt in zip(t.rows[0].cells, ("统计项", "数量")):
        _set_font(c.paragraphs[0].add_run(txt), size=10, bold=True)
    for k, v in stats.get("all", {}).items():
        row = t.add_row().cells
        _set_font(row[0].paragraphs[0].add_run(str(k)), size=9)
        _set_font(row[1].paragraphs[0].add_run(str(v)), size=9)

    mode = stats.get("mode", "set")
    if is_compare and mode == "excelmpn":
        doc.add_heading("Excel MPN 对比明细", level=1)
        headers = ["MPN(物料号)", "结果", "数量A", "数量B", "位号A", "位号B", "厂商A", "厂商B", "差异明细"]
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        for c, txt in zip(table.rows[0].cells, headers):
            _set_font(c.paragraphs[0].add_run(txt), size=9, bold=True)
        RED = (0x9C, 0x00, 0x06)
        for row in rows:
            changed = row.get("status") != "一致"
            ra = row.get("rec_a") or {}
            rb = row.get("rec_b") or {}
            desc = "；".join("%s: %s→%s" % (d["field"], d["a"] or "空", d["b"] or "空")
                             for d in row["field_diffs"])
            cells = table.add_row().cells
            vals = [row["item"], row["status"], ra.get("qty", ""), rb.get("qty", ""),
                    ",".join(ra.get("refs", [])), ",".join(rb.get("refs", [])),
                    ra.get("mfg", ""), rb.get("mfg", ""), desc]
            for c, v in zip(cells, vals):
                _set_font(c.paragraphs[0].add_run(str(v)), size=9, color=(RED if changed else None))
    elif is_compare and mode == "mpn":
        doc.add_heading("MPN 对比明细", level=1)
        headers = ["MPN(料号)", "结果", "Excel数量", "PDF数量", "Excel位号", "PDF位号", "厂商", "PDF页"]
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        for c, txt in zip(table.rows[0].cells, headers):
            _set_font(c.paragraphs[0].add_run(txt), size=9, bold=True)
        RED = (0x9C, 0x00, 0x06)
        for row in rows:
            changed = row.get("status") != "一致"
            cells = table.add_row().cells
            vals = [row.get("item", ""), row.get("status", ""), row.get("qty_a", ""),
                    row.get("qty_b", ""), row.get("refs_a", ""), row.get("refs_b", ""),
                    row.get("mfg", ""), row.get("pages", "")]
            for c, v in zip(cells, vals):
                _set_font(c.paragraphs[0].add_run(str(v)), size=9, color=(RED if changed else None))
    elif is_compare:
        doc.add_heading("对比明细", level=1)
        headers = ["器件/位号", "结果", "文件A", "文件B"]
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        for c, txt in zip(table.rows[0].cells, headers):
            _set_font(c.paragraphs[0].add_run(txt), size=9, bold=True)
        RED = (0x9C, 0x00, 0x06)
        for row in rows:
            changed = row.get("status") != "一致"
            cells = table.add_row().cells
            vals = [row.get("item", ""), row.get("status", ""),
                    row.get("a_pages") or row.get("a", ""), row.get("b_pages") or row.get("b", "")]
            for c, v in zip(cells, vals):
                _set_font(c.paragraphs[0].add_run(str(v)), size=9, color=(RED if changed else None))
    else:
        doc.add_heading("核对明细", level=1)
        headers = ["BOM行号", "列", "原值", "位号/值", "PDF中存在", "所在页码", "结果"]
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        for c, txt in zip(table.rows[0].cells, headers):
            _set_font(c.paragraphs[0].add_run(txt), size=9, bold=True)
        RED = (0x9C, 0x00, 0x06)
        for row in rows:
            for d in row.get("designators", []):
                cells = table.add_row().cells
                vals = [row["header_row"], row["col_name"], row["raw"], d["designator"],
                        "是" if d["found"] else "否",
                        ",".join(str(p) for p in d["pages"]) if d["found"] else "",
                        "一致" if d["found"] else "不一致"]
                for c, v in zip(cells, vals):
                    _set_font(c.paragraphs[0].add_run(str(v)), size=9,
                              color=(RED if not d["found"] else None))
    doc.save(docx_path)
    return docx_path


# ============================================================
#  美化 GUI（含二级子窗口）
# ============================================================
import tkinter as tk
from tkinter import ttk, filedialog, messagebox

# ---------- 配置持久化（记住上次选择） ----------------------
def _cfg_path():
    return os.path.join(os.path.expanduser("~"), ".bom_verify_config.json")


def _load_config():
    try:
        with open(_cfg_path(), "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return {}


def _save_config(cfg):
    try:
        with open(_cfg_path(), "w", encoding="utf-8") as f:
            json.dump(cfg, f, ensure_ascii=False, indent=2)
    except Exception:
        pass


def _smart_dir(app):
    """智能默认目录：优先输出文件夹→配置记录→Bom-data→脚本目录。"""
    for cand in (app.get("out_dir"), (_load_config() or {}).get("out_dir"),
                 r"D:\Evan\Bom-data"):
        if cand and os.path.isdir(cand):
            return cand
    return os.path.dirname(os.path.abspath(__file__))

# ---------- 配色主题 --------------------------------
class Theme:
    BG        = "#F4F6FB"   # 主背景色
    CARD      = "#FFFFFF"   # 卡片底色
    BRAND     = "#2C6FBB"   # 顶部横幅+按钮 主色调（改成你喜欢的主色）
    BRAND_DK  = "#1E4E84"   #   主色加深
    BRAND_LT  = "#E8F0FA"   #   主色浅底（表头用）
    TEXT      = "#23303B"   # 正文文字色
    SUBTEXT   = "#6B7A8D"   # 次要文字色
    LINE      = "#DFE6F0"   # 边框/分隔线色
    OK_BG     = "#E7F7EA"   # 一致行的绿底
    OK_FG     = "#1E7B34"   # 一致行的绿字
    BAD_BG    = "#FDEAEA"   # 不一致行的红底
    BAD_FG    = "#C0392B"   # 不一致行的红字
    WARN_BG   = "#FFF4E0"   # 警告行的黄底
    WARN_FG   = "#996A00"   # 警告行的黄字
    BUTTON_OK = "#2E8B57"


# ---------- Windows 高DPI 适配 ----------
# 解决 tkinter 在 Windows 高分屏下文字模糊、发虚的问题。
def enable_dpi_awareness():
    """在创建任何窗口前调用，让程序感知 Windows 缩放比例(DPI)。"""
    if sys.platform == "win32":
        try:
            import ctypes
            # 0=无效, 1=系统DPI感知, 2=每显示器DPI感知(最清晰)
            ctypes.windll.shcore.SetProcessDpiAwareness(2)
        except Exception:
            try:
                import ctypes
                ctypes.windll.user32.SetProcessDPIAware()
            except Exception:
                pass


def detect_scaling(root):
    """根据屏幕 DPI 计算字体缩放系数，返回 (scale, 单位字体尺寸)。"""
    try:
        import ctypes
        dpi = float(ctypes.windll.shcore.GetScaleFactorForDevice(0))
    except Exception:
        dpi = 96.0
    scale = dpi / 96.0
    # 用 tkinter 的 scaling 保持字体物理尺寸一致
    try:
        root.tk.call("tk", "scaling", scale)
    except Exception:
        pass
    return scale


# 统一字体规格(避免大小/风格不一致)
def _f(*size_weight):
    """生成统一字体元组，默认用 Segoe UI（Windows 自带现代字体，清晰圆润）。
    用法: _f(10)   → 常规10号;   _f(10, "bold") → 粗体10号
    """
    size = size_weight[0]
    weight = size_weight[1] if len(size_weight) > 1 else "normal"
    return ("Segoe UI", int(size), weight)


def _style(theme, scale=1.0):
    s = ttk.Style()
    try:
        s.theme_use("clam")
    except Exception:
        pass
    base = max(10, round(10 * scale))
    s.configure(".", background=theme.BG, foreground=theme.TEXT, font=_f(base))
    s.configure("Card.TFrame", background=theme.CARD)
    s.configure("Card.TLabelframe", background=theme.CARD, bordercolor=theme.LINE,
                relief="solid", borderwidth=1)
    s.configure("Card.TLabelframe.Label", background=theme.CARD, foreground=theme.BRAND_DK,
                font=_f(base, "bold"))
    s.configure("TLabel", background=theme.BG, foreground=theme.TEXT, font=_f(base))
    s.configure("CLabel.TLabel", background=theme.CARD, font=_f(base))
    s.configure("Sub.TLabel", background=theme.CARD, foreground=theme.SUBTEXT,
                font=_f(max(9, base - 1)))
    s.configure("Title.TLabel", background=theme.CARD, foreground=theme.BRAND_DK,
                font=_f(max(15, round(base * 1.5)), "bold"))
    s.configure("Count.TLabel", background=theme.CARD, font=_f(max(11, base), "bold"))
    s.configure("OK.TLabel", background=theme.CARD, foreground=theme.OK_FG, font=_f(base, "bold"))
    s.configure("BAD.TLabel", background=theme.CARD, foreground=theme.BAD_FG, font=_f(base, "bold"))

    s.configure("TButton", background="#E7EDF6", foreground=theme.TEXT, bordercolor=theme.LINE,
                padding=(12, 6), font=_f(base))
    s.map("TButton", background=[("active", "#D5E2F4")])
    s.configure("Primary.TButton", background=theme.BRAND, foreground="#FFFFFF", bordercolor=theme.BRAND,
                padding=(16, 8), font=_f(base, "bold"))
    s.map("Primary.TButton", background=[("active", theme.BRAND_DK)])
    s.configure("Danger.TButton", background="#E8B4B4", foreground="#8B0000", bordercolor="#C88080",
                padding=(14, 6), font=_f(base))
    s.map("Danger.TButton", background=[("active", "#DCA0A0")])
    s.configure("Success.TButton", background="#5B9E6B", foreground="#FFF", bordercolor="#4A8B58",
                padding=(14, 6), font=_f(base))
    s.map("Success.TButton", background=[("active", "#478457")])

    s.configure("TEntry", fieldbackground="#FFF", foreground=theme.TEXT, bordercolor=theme.LINE,
                font=_f(base))
    s.configure("TCombobox", fieldbackground="#FFF", foreground=theme.TEXT,
                bordercolor=theme.LINE, arrowcolor=theme.BRAND, font=_f(base))

    s.configure("Treeview", background="#FFFFFF", fieldbackground="#FFFFFF", foreground=theme.TEXT,
                rowheight=max(26, round(26 * scale)), bordercolor=theme.LINE, font=_f(base))
    s.configure("Treeview.Heading", background=theme.BRAND_LT, foreground=theme.BRAND_DK,
                font=_f(base, "bold"), relief="flat")
    s.map("Treeview", background=[("selected", "#BFD8F2")])
    s.configure("Vertical.TScrollbar", background="#DAE4F0", troughcolor=theme.BG, arrowcolor=theme.BRAND_DK)
    s.configure("Horizontal.TScrollbar", background="#DAE4F0", troughcolor=theme.BG, arrowcolor=theme.BRAND_DK)


def _make_card(parent, title, padding=(14, 12, 14, 12)):
    lf = ttk.LabelFrame(parent, text=title, style="Card.TLabelframe", padding=padding)
    return lf


def _icon_text(label, color, bg):
    w = tk.Label(label.master, text=" ", bg=bg, fg=color, padx=4)
    return w


# ---------- 二级窗口：PDF 文字查看 --------------------
def open_pdf_text_win(root, app):
    from tkinter import Text
    win = tk.Toplevel(root)
    win.title("PDF 原理图解码文字")
    win.geometry("760x560")
    win.configure(bg=Theme.BG)
    win.transient(root)
    win.grab_set()

    head = ttk.Frame(win, style="Card.TFrame", padding=10)
    head.pack(fill="x", padx=10, pady=(10, 6))
    ttk.Label(head, text="PDF 每页解码出的文字（含位置坐标）",
              style="Title.TLabel").pack(side="left")
    ttk.Label(head, text="", style="Sub.TLabel").pack(side="left")

    ttk.Style().configure("page.Tab", padding=(20, 6))

    pw = ttk.Panedwindow(win, orient="horizontal")
    pw.pack(fill="both", expand=True, padx=10, pady=(0, 10))

    left = ttk.Frame(pw)
    pw.add(left, weight=0)
    ttk.Label(left, text="选择页码:", style="Sub.TLabel").pack(pady=(4, 2))
    lb = tk.Listbox(left, width=10, height=28, activestyle="dotbox",
                    selectbackground=Theme.BRAND, selectforeground="#fff",
                    font=_f(10), bg="#FFF", fg=Theme.TEXT, relief="flat",
                    highlightbackground=Theme.LINE, highlightthickness=1)
    lb.pack(fill="y", expand=True)

    right = ttk.Frame(pw)
    pw.add(right, weight=1)
    txt = Text(right, wrap="none", font=_f(10),
               bg="#FFFFFF", fg=Theme.TEXT, insertbackground=Theme.BRAND,
               relief="flat", padx=10, pady=8)
    vs = ttk.Scrollbar(right, orient="vertical", command=txt.yview)
    hs = ttk.Scrollbar(right, orient="horizontal", command=txt.xview)
    txt.configure(yscrollcommand=vs.set, xscrollcommand=hs.set)
    txt.grid(row=0, column=0, sticky="nsew")
    vs.grid(row=0, column=1, sticky="ns")
    hs.grid(row=1, column=0, sticky="ew")
    right.rowconfigure(0, weight=1)
    right.columnconfigure(0, weight=1)

    pages_store = {}
    for pno in sorted(app.get("pdf_pages_text", {})):
        lb.insert("end", f"第 {pno} 页")
        pl = app["pdf_pages_text"][pno]
        body = []
        for y, line in pl:
            body.append(f"{y:<8.1f}  {line}")
        pages_store[pno] = "\n".join(body)

    def show(_=None):
        sel = lb.curselection()
        if not sel:
            return
        pno = int(lb.get(sel[0]).replace("第 ", "").replace(" 页", ""))
        txt.delete("1.0", "end")
        txt.insert("1.0", pages_store.get(pno, "(无文字)"))

    lb.bind("<<ListboxSelect>>", show)
    if lb.size():
        lb.selection_set(0)
        show()


# ---------- 二级窗口：位号详情 ------------------------
def open_detail_win(root, app, designator, pages, row, col, raw):
    win = tk.Toplevel(root)
    win.title(f"位号详情  {designator}")
    win.geometry("640x420")
    win.configure(bg=Theme.BG)
    win.transient(root)
    win.grab_set()

    card = _make_card(win, "查询信息")
    card.pack(fill="x", padx=12, pady=10)
    grid = ttk.Frame(card, style="Card.TFrame")
    grid.pack(fill="x")
    rows = [
        ("位号", designator),
        ("BOM行", str(row)),
        ("核对列", col),
        ("原值(BOM)", str(raw)),
        ("PDF页码", ", ".join(str(p) for p in pages) if pages else "未找到"),
        ("结果", "一致" if pages else "不一致"),
    ]
    for i, (k, v) in enumerate(rows):
        ttk.Label(grid, text=k, style="Sub.TLabel", width=10).grid(row=i, column=0, sticky="ne", pady=2)
        ttk.Label(grid, text=v, style="CLabel.TLabel", anchor="w").grid(row=i, column=1, sticky="w", padx=8, pady=2)
        if k == "结果":
            if pages:
                ttk.Label(grid, text="一致", style="OK.TLabel").grid(row=i, column=2, sticky="w")
            else:
                ttk.Label(grid, text="不一致", style="BAD.TLabel").grid(row=i, column=2, sticky="w")

    ctx = _make_card(win, "位号在 PDF 中出现的行（上下文）")
    ctx.pack(fill="both", expand=True, padx=12, pady=(0, 12))
    area = tk.Text(ctx, height=12, font=_f(10), bg="#FFF",
                   fg=Theme.TEXT, relief="flat", padx=8, pady=6)
    area.pack(side="left", fill="both", expand=True)
    sv = ttk.Scrollbar(ctx, orient="vertical", command=area.yview)
    sv.pack(side="right", fill="y")
    area.configure(yscrollcommand=sv.set)
    area.insert("1.0", "（未找到该位号在 PDF 中的上下文）" if not pages else "")
    if pages:
        # 显示每页包含该位号的行
        de = app.get("designator_context", {})
        lines = de.get(designator, [])
        area.delete("1.0", "end")
        if lines:
            for ln in lines:
                area.insert("end", ln + "\n")
        else:
            area.insert("1.0", "该位号已找到（页码见上）。可打开“PDF文字”窗口查看上下文。")


# ---------- 二级窗口：不一致清单 ----------------------
def open_mismatch_win(root, app):
    stats = app.get("stats")
    if not stats:
        messagebox.showinfo("提示", "请先执行“开始核对”。")
        return
    win = tk.Toplevel(root)
    win.title("不一致清单")
    win.geometry("760x560")
    win.configure(bg=Theme.BG)
    win.transient(root)

    card = _make_card(win, "BOM 有、PDF 无（不一致）")
    card.pack(fill="both", expand=True, padx=12, pady=8)
    tree = ttk.Treeview(card, columns=("des", "row", "raw"), show="headings", height=12)
    tree.heading("des", text="位号")
    tree.heading("row", text="BOM 行")
    tree.heading("raw", text="原值")
    tree.column("des", width=140)
    tree.column("row", width=80)
    tree.column("raw", width=420)
    vs = ttk.Scrollbar(card, orient="vertical", command=tree.yview)
    tree.configure(yscrollcommand=vs.set)
    tree.pack(side="left", fill="both", expand=True)
    vs.pack(side="right", fill="y")
    for b in stats["bad"]:
        tree.insert("", "end", values=(b["designator"], b.get("row", ""), ""), tags="bad")
    tree.tag_configure("bad", background=Theme.BAD_BG, foreground=Theme.BAD_FG)

    card2 = _make_card(win, "PDF 有、BOM 无（额外器件/网络）")
    card2.pack(fill="both", expand=True, padx=12, pady=(0, 12))
    tree2 = ttk.Treeview(card2, columns=("des", "pages"), show="headings", height=10)
    tree2.heading("des", text="位号")
    tree2.heading("pages", text="PDF 页码")
    tree2.column("des", width=140)
    tree2.column("pages", width=160)
    vs2 = ttk.Scrollbar(card2, orient="vertical", command=tree2.yview)
    tree2.configure(yscrollcommand=vs2.set)
    tree2.pack(side="left", fill="both", expand=True)
    vs2.pack(side="right", fill="y")
    for e in stats.get("extra", []):
        tree2.insert("", "end", values=(e, ",".join(str(p) for p in app.get("pdf_tokens", {}).get(e, []))),
                     tags="warn")
    tree2.tag_configure("warn", background=Theme.WARN_BG, foreground=Theme.WARN_FG)


# ---------- 二级窗口：汇总统计 ------------------------
def open_summary_win(root, app):
    stats = app.get("stats")
    if not stats:
        messagebox.showinfo("提示", "请先执行“开始核对”。")
        return
    win = tk.Toplevel(root)
    win.title("核对汇总")
    win.geometry("460x300")
    win.configure(bg=Theme.BG)
    win.transient(root)
    card = _make_card(win, "统计")
    card.pack(fill="both", expand=True, padx=12, pady=12)
    for k, v in stats["all"].items():
        row = ttk.Frame(card, style="Card.TFrame")
        row.pack(fill="x", pady=3)
        ttk.Label(row, text=k, style="CLabel.TLabel", width=32, anchor="w").pack(side="left")
        ttk.Label(row, text=str(v), style="Count.TLabel", anchor="e").pack(side="right")


# ============================================================
#  主窗口
# ============================================================
def run_gui():
    import tkinter as tk
    from tkinter import ttk

    enable_dpi_awareness()          # ★ 必须在创建窗口前调用：解决模糊

    root = tk.Tk()
    root.title("BOM ↔ PDF 原理图 器件核对工具 V1")
    scale = detect_scaling(root)    # 依据屏幕 DPI 计算缩放
    _style(Theme, scale)            # 字体随缩放联动，避免大小不一致

    # 自适应窗口：以屏幕尺寸为上限，防止小屏幕"显示不全"
    sw, sh = root.winfo_screenwidth(), root.winfo_screenheight()
    W = int(1280 * scale)
    H = int(860 * scale)
    if W > sw - 40:
        W = sw - 40
    if H > sh - 80:
        H = sh - 80
    root.geometry(f"{W}x{H}")
    root.minsize(int(940 * scale) if scale > 1 else 940,
                 int(640 * scale) if scale > 1 else 640)
    root.configure(bg=Theme.BG, padx=8, pady=8)

    # 全局数据
    app = {
        "headers": [], "data": [], "bom_path": "", "pdf_path": "",
        "pdf_tokens": {}, "pdf_all_words": [], "pdf_pages_text": {},
        "result_rows": [], "stats": None, "designator_context": {},
        "mode": "BOM vs PDF",           # 当前对比类型
        "file_b": "",                   # 第二个文件（对比用）
        "headers_b": [], "data_b": [],
    }
    lb_xlsx = tk.StringVar()
    lb_pdf = tk.StringVar()
    lb_file_b = tk.StringVar()
    lb_compare = tk.StringVar(value="BOM vs PDF")
    lb_col = tk.StringVar()
    lb_mode = tk.StringVar(value="位号核对(推荐)")
    var_status = tk.StringVar(value="就绪：请选择对比类型与文件")

    # ---------- 恢复上次配置 ----------
    cfg = _load_config()
    if cfg.get("compare"):
        lb_compare.set(cfg["compare"])
    if cfg.get("col"):
        lb_col.set(cfg["col"])
    if cfg.get("mode"):
        lb_mode.set(cfg["mode"])
    for key, var, apk in (("bom", lb_xlsx, "bom_path"),
                          ("pdf", lb_pdf, "pdf_path"),
                          ("file_b", lb_file_b, "file_b")):
        p = cfg.get(key)
        if p and os.path.exists(p):
            var.set(os.path.basename(p))
            app[apk] = p
    if cfg.get("out_dir") and os.path.isdir(cfg["out_dir"]):
        app["out_dir"] = cfg["out_dir"]
    if "adv_settings" in cfg:
        app["adv_settings"] = cfg["adv_settings"]

    # ---------- 顶部标题栏 ----------
    header = tk.Frame(root, bg=Theme.BRAND)
    header.pack(fill="x", pady=(0, 8))
    tk.Label(header, text="BOM ↔ PDF 器件核对与对比工具", bg=Theme.BRAND, fg="#FFFFFF",
             font=_f(max(15, round(15 * scale)), "bold")).pack(side="left", padx=14, pady=10)
    tk.Label(header, text="三种对比方案：Excel-Excel(MPN主键) · Excel-PDF · PDF-PDF",
             bg=Theme.BRAND, fg="#DCE8FB", font=_f(max(9, round(9 * scale)))).pack(side="left", padx=6)
    tk.Button(header, text="复位配置", bg=Theme.BRAND_DK, fg="#FFFFFF", relief="flat",
              activebackground=Theme.BRAND_DK, activeforeground="#FFFFFF", bd=0,
              font=_f(max(9, round(9 * scale))), padx=10, pady=2,
              command=lambda: (_save_config({}), root.after(0, lambda: messagebox.showinfo(
                  "复位配置", "已清空上次记忆，重启程序生效。", parent=root)))).pack(side="right", padx=12)

    # ---------- 主内容区 ----------
    main = ttk.Frame(root)
    main.pack(fill="both", expand=True)

    # 左栏：放入可滚动容器（小屏也能全部看到，解决"显示不全"）
    left_shell = ttk.Frame(main)
    left_shell.pack(side="left", fill="y", padx=(0, 8))
    left_canvas = tk.Canvas(left_shell, bg=Theme.BG, highlightthickness=0, width=int(440 * scale))
    left_vsb = ttk.Scrollbar(left_shell, orient="vertical", command=left_canvas.yview)
    left_canvas.configure(yscrollcommand=left_vsb.set)
    left_canvas.pack(side="left", fill="both", expand=True)
    left_vsb.pack(side="right", fill="y")

    left_col = ttk.Frame(left_canvas)
    left_win = left_canvas.create_window((0, 0), window=left_col, anchor="nw")
    left_col.bind("<Configure>", lambda e: left_canvas.configure(scrollregion=left_canvas.bbox("all")))
    left_canvas.bind("<Configure>", lambda e: left_canvas.itemconfigure(left_win, width=e.width))

    def _on_left_wheel(event):
        left_canvas.yview_scroll(int(-event.delta / 120), "units")
    left_canvas.bind("<Enter>", lambda e: left_canvas.bind_all("<MouseWheel>", _on_left_wheel))
    left_canvas.bind("<Leave>", lambda e: left_canvas.unbind_all("<MouseWheel>"))

    card_compare = _make_card(left_col, "1 · 对比类型")
    card_compare.pack(fill="x", pady=(0, 8))
    _compare_type_row(card_compare, lb_compare, lb_mode, lb_col, lb_xlsx, lb_pdf, lb_file_b,
                      app, var_status)

    card_file = _make_card(left_col, "2 · 选择文件")
    card_file.pack(fill="x", pady=(0, 8))
    _file_rows(card_file, lb_compare, lb_xlsx, lb_pdf, lb_file_b, app, var_status, lb_col, root)

    card_key = _make_card(left_col, "3 · 核对设置")
    card_key.pack(fill="x", pady=(0, 8))
    _setting_rows(card_key, lb_col, lb_mode, app, lb_xlsx, lb_pdf, lb_file_b, lb_compare,
                  var_status, root)

    card_prob = _make_card(left_col, "4 · 结果窗口（二级）")
    card_prob.pack(fill="x")
    _subwin_rows(card_prob, root, app)

    # 右栏：预览与结果（自适应拉伸）
    right_col = ttk.Frame(main)
    right_col.pack(side="left", fill="both", expand=True)
    right_col.grid_rowconfigure(1, weight=3)
    right_col.grid_rowconfigure(0, weight=1)
    right_col.grid_columnconfigure(0, weight=1)

    _preview_card(right_col, app, lb_col)
    _result_card(right_col, root, app)

    # ---------- 底部状态栏 ----------
    status_bar = tk.Frame(root, bg=Theme.CARD, highlightbackground=Theme.LINE, highlightthickness=1)
    status_bar.pack(fill="x", pady=(8, 0))
    tk.Label(status_bar, textvariable=var_status, bg=Theme.CARD, fg=Theme.SUBTEXT,
             font=_f(max(9, round(9 * scale))), anchor="w", padx=10, pady=6).pack(side="left", fill="x", expand=True)
    # 进度条（核对/解码时显示）
    pb = ttk.Progressbar(status_bar, mode="indeterminate", length=140, maximum=100)
    pb.pack(side="right", padx=10, pady=4)
    app["progressbar"] = pb

    # ---------- 关闭时保存配置 ----------
    def _on_close():
        _save_config({
            "compare": lb_compare.get(),
            "mode": lb_mode.get(),
            "col": lb_col.get(),
            "bom": app.get("bom_path") or lb_xlsx.get(),
            "pdf": app.get("pdf_path") or lb_pdf.get(),
            "file_b": app.get("file_b") or lb_file_b.get(),
            "out_dir": app.get("out_dir") or "",
            "adv_settings": bool(app.get("adv_settings", False)),
        })
        root.destroy()
    root.protocol("WM_DELETE_WINDOW", _on_close)

    # ---------- 自动载入同目录文件 ----------
    root.after(200, lambda: _auto_load(app, lb_xlsx, lb_pdf, lb_col))
    root.mainloop()


# ---------- 对比类型行 ----------
def _compare_type_row(parent, lb_compare, lb_mode, lb_col, lb_xlsx, lb_pdf, lb_file_b,
                      app, var_status):
    row = ttk.Frame(parent, style="Card.TFrame")
    row.pack(fill="x", pady=3)
    ttk.Label(row, text="对比类型", style="Sub.TLabel", width=11, anchor="w").pack(side="left")
    cmb = ttk.Combobox(row, textvariable=lb_compare, width=26, state="readonly")
    cmb.pack(side="left", fill="x", expand=True)
    cmb["values"] = ["Excel vs Excel", "Excel vs PDF", "PDF vs PDF"]
    if not lb_compare.get():
        cmb.current(0)
    cmb.bind("<<ComboboxSelected>>", lambda e: _on_compare_change(
        lb_compare, lb_mode, lb_col, lb_xlsx, lb_pdf, lb_file_b, app, var_status))


def _on_compare_change(lb_compare, lb_mode, lb_col, lb_xlsx, lb_pdf, lb_file_b, app, var_status):
    t = lb_compare.get()
    app["mode"] = t
    if t == "Excel vs Excel":
        lb_mode.set("MPN对比")
        var_status.set("方案三：以 MPN(物料号) 为主键比对两个 Excel 类文件（CSV/xlsx/xls/xlsm）")
    elif t == "PDF vs PDF":
        lb_mode.set("位号对比")
        var_status.set("方案一：比对两个 PDF 原理图的位号")
    else:
        lb_mode.set("位号核对(推荐)")
        var_status.set("方案二：Excel 类文件(BOM) 与 PDF 原理图 器件核对")


# ---------- 文件行（三模式动态） ----------
def _file_rows(parent, lb_compare, lb_xlsx, lb_pdf, lb_file_b, app, var_status, lb_col, root):
    # 行 A（文件 A）：Excel vs Excel 时=Excel A；Excel vs PDF 时=Excel BOM；PDF vs PDF 时=PDF A
    rA = ttk.Frame(parent, style="Card.TFrame")
    rA.pack(fill="x", pady=3)
    lab_a = ttk.Label(rA, text="文件 A", style="Sub.TLabel", width=11, anchor="w")
    lab_a.pack(side="left")
    eA = ttk.Entry(rA, textvariable=lb_xlsx, width=30)
    eA.pack(side="left", fill="x", expand=True)
    ttk.Button(rA, text="浏览", width=6,
               command=lambda: _browse_file(app, "A", lb_compare, lb_xlsx, lb_pdf, lb_file_b,
                                            var_status, lb_col, root)).pack(side="left", padx=(4, 0))

    def _bind_a():
        if lb_compare.get() == "PDF vs PDF":
            eA.configure(textvariable=lb_pdf)
        else:
            eA.configure(textvariable=lb_xlsx)

    # 行 B（文件 B）：Excel vs Excel 时=Excel B；其余模式=PDF B
    rB = ttk.Frame(parent, style="Card.TFrame")
    rB.pack(fill="x", pady=3)
    lab_b = ttk.Label(rB, text="文件 B", style="Sub.TLabel", width=11, anchor="w")
    lab_b.pack(side="left")
    eB = ttk.Entry(rB, textvariable=lb_pdf, width=30)
    eB.pack(side="left", fill="x", expand=True)
    ttk.Button(rB, text="浏览", width=6,
               command=lambda: _browse_file(app, "B", lb_compare, lb_xlsx, lb_pdf, lb_file_b,
                                            var_status, lb_col, root)).pack(side="left", padx=(4, 0))

    def _bind_b():
        if lb_compare.get() == "Excel vs Excel":
            eB.configure(textvariable=lb_file_b)
        else:
            eB.configure(textvariable=lb_pdf)

    # 行 C（占位，恒隐藏：文件 B 统一走 rB）
    rC = ttk.Frame(parent, style="Card.TFrame")

    def _refresh():
        t = lb_compare.get()
        _bind_a()
        _bind_b()
        if t == "Excel vs Excel":
            lab_a.configure(text="文件 A (Excel)")
            lab_b.configure(text="文件 B (Excel)")
            # 行B 用 lb_file_b 显示第二个 Excel；隐藏行C
            rC.pack_forget()
            rB.pack(in_=parent, fill="x", pady=3)
        elif t == "PDF vs PDF":
            lab_a.configure(text="文件 A (PDF)")
            lab_b.configure(text="文件 B (PDF)")
            # 两个 PDF：行B 绑定 lb_file_b 显示第二个 PDF
            rC.pack_forget()
            rB.pack(in_=parent, fill="x", pady=3)
        else:
            lab_a.configure(text="Excel BOM")
            lab_b.configure(text="原理图 PDF")
            rC.pack_forget()
            rB.pack(in_=parent, fill="x", pady=3)

    lb_compare.trace("w", lambda *a: _refresh())
    parent._refresh_files = _refresh
    _refresh()


def _browse_file(app, slot, lb_compare, lb_xlsx, lb_pdf, lb_file_b, var_status, lb_col, root):
    t = lb_compare.get()
    d = _smart_dir(app)
    if t == "Excel vs Excel":
        p = filedialog.askopenfilename(title="选择表格文件", initialdir=d,
                                       filetypes=[("表格文件", "*.xlsx *.xls *.xlsb *.ods *.csv *.tsv *.txt *.html *.htm"),
                                                  ("所有文件", "*.*")])
        if not p:
            return
        if slot == "A":
            app["bom_path"] = p
            lb_xlsx.set(os.path.basename(p))
            try:
                headers, data = load_bom(p)
                app["headers"], app["data"] = headers, data
                lb_col.set(_designator_col_name(headers) or (headers[0] if headers else ""))
                _refresh_preview(app, lb_col)
            except Exception as e:
                messagebox.showerror("错误", "Excel A 读取失败:\n%s" % e, parent=root)
                return
        else:
            app["file_b"] = p
            lb_file_b.set(os.path.basename(p))
            try:
                headers_b, data_b = load_bom(p)
                app["headers_b"], app["data_b"] = headers_b, data_b
                # 通知核对设置刷新 文件B主键列/数量列 下拉
                for cb in (app.get("_col_refresh_callbacks") or []):
                    try:
                        cb()
                    except Exception:
                        pass
            except Exception as e:
                messagebox.showerror("错误", "Excel B 读取失败:\n%s" % e, parent=root)
                return
        var_status.set("已选 Excel：A=%s，B=%s" % (os.path.basename(app["bom_path"]),
                                                os.path.basename(app["file_b"])))
        return

    if t == "PDF vs PDF":
        p = filedialog.askopenfilename(title="选择 PDF", initialdir=d,
                                       filetypes=[("PDF", "*.pdf"), ("所有文件", "*.*")])
        if not p:
            return
        if slot == "A":
            app["pdf_path"] = p
            lb_pdf.set(os.path.basename(p))
        else:
            app["file_b"] = p
            lb_file_b.set(os.path.basename(p))
        var_status.set("已选 PDF：A=%s，B=%s" % (os.path.basename(app["pdf_path"]),
                                                os.path.basename(app["file_b"])))
        return

    # Excel vs PDF：A=Excel BOM，B=原理图 PDF
    if slot == "A":
        _browse("BOM", app, root, var_status, lb_col)
    else:
        _browse("PDF", app, root, var_status, lb_col)


def _browse(kind, app, root, var_status, lb_col):
    d = _smart_dir(app)
    if kind == "BOM":
        p = filedialog.askopenfilename(title="选择 BOM 表格", initialdir=d,
                                       filetypes=[("表格文件", "*.xlsx *.xls *.xlsb *.ods *.csv *.tsv *.txt *.html *.htm"),
                                                  ("所有文件", "*.*")])
        if p:
            app["bom_path"] = p
            try:
                headers, data = load_bom(p)
                app["headers"], app["data"] = headers, data
                if lb_col is not None:
                    lb_col.set(_designator_col_name(headers) or (headers[0] if headers else ""))
                _refresh_preview(app, lb_col)
                var_status.set("已载入 BOM，共 %d 行数据" % len(data))
            except Exception as e:
                messagebox.showerror("错误", "BOM 读取失败:\n%s" % e, parent=root)
        return
    p = filedialog.askopenfilename(title="选择原理图 PDF", initialdir=d,
                                   filetypes=[("PDF", "*.pdf"), ("所有文件", "*.*")])
    if p:
        app["pdf_path"] = p
        var_status.set("已选择 PDF：%s" % os.path.basename(p))


def _refresh_preview(app, lb_col):
    tree = app.get("preview_tree")
    if not tree or not app["headers"]:
        return
    tree["columns"] = app["headers"]
    for h in app["headers"]:
        tree.heading(h, text=h)
        tree.column(h, width=100, anchor="w")
    for i in tree.get_children():
        tree.delete(i)
    for row in app["data"][:10]:
        tree.insert("", "end", values=[str(v)[:60] for v in row["values"]])
    if lb_col is not None and app["headers"]:
        lb_col.set(_designator_col_name(app["headers"]) or app["headers"][0])
    # 通知"核对设置"刷新列下拉（核对列 / Excel列）
    for cb in (app.get("_col_refresh_callbacks") or []):
        try:
            cb()
        except Exception:
            pass


# ---------- 核对设置 ----------
def _setting_rows(parent, lb_col, lb_mode, app, lb_xlsx, lb_pdf, lb_file_b, lb_compare,
                  var_status, root):
    # 输出文件夹：用户指定报告导出路径
    r0 = ttk.Frame(parent, style="Card.TFrame")
    r0.pack(fill="x", pady=3)
    ttk.Label(r0, text="输出文件夹", style="Sub.TLabel", width=11, anchor="w").pack(side="left")
    var_out = tk.StringVar(value=app.get("out_dir") or "")
    e_out = ttk.Entry(r0, textvariable=var_out, width=28)
    e_out.pack(side="left", fill="x", expand=True)
    ttk.Button(r0, text="…", width=3, style="TButton",
               command=lambda: _pick_out_dir(app, var_out, root)).pack(side="left", padx=(4, 0))

    # 导出文件类型选择：Excel / Word / PDF / TXT
    rFmt = ttk.Frame(parent, style="Card.TFrame")
    rFmt.pack(fill="x", pady=3)
    ttk.Label(rFmt, text="导出类型", style="Sub.TLabel", width=11, anchor="w").pack(side="left")
    default_fmt = app.get("export_formats") or ["xlsx", "docx", "pdf", "txt"]
    var_xlsx = tk.BooleanVar(value="xlsx" in default_fmt)
    var_docx = tk.BooleanVar(value="docx" in default_fmt)
    var_pdf = tk.BooleanVar(value="pdf" in default_fmt)
    var_txt = tk.BooleanVar(value="txt" in default_fmt)
    ttk.Checkbutton(rFmt, text="Excel", variable=var_xlsx, style="TCheckbutton").pack(side="left", padx=(0, 6))
    ttk.Checkbutton(rFmt, text="Word", variable=var_docx, style="TCheckbutton").pack(side="left", padx=(0, 6))
    ttk.Checkbutton(rFmt, text="PDF", variable=var_pdf, style="TCheckbutton").pack(side="left", padx=(0, 6))
    ttk.Checkbutton(rFmt, text="TXT", variable=var_txt, style="TCheckbutton").pack(side="left")
    app["_fmt_vars"] = (var_xlsx, var_docx, var_pdf, var_txt)

    # 高级比对设置：手动选择是否展开
    rTop = ttk.Frame(parent, style="Card.TFrame")
    rTop.pack(fill="x", pady=(6, 2))
    var_adv = tk.BooleanVar(value=bool(app.get("adv_settings", False)))
    ttk.Checkbutton(rTop, text="显示高级比对设置（主键列/数量列等）", variable=var_adv,
                    style="TCheckbutton").pack(side="left")

    # 文件A / 文件B 主键列：支持多选（Ctrl/Shift 多选；不选则自动识别 MPN/料号列）
    rKA = ttk.Frame(parent, style="Card.TFrame")
    ttk.Label(rKA, text="文件A主键列", style="Sub.TLabel", width=11, anchor="w").pack(side="top", anchor="w")
    var_key_a = tk.StringVar(value=",".join(app.get("key_a", []) or []))
    e_key_a = ttk.Entry(rKA, textvariable=var_key_a, width=30)
    e_key_a.pack(side="top", fill="x", pady=(0, 2))
    lbx_a = tk.Listbox(rKA, width=30, height=5, exportselection=False,
                       selectmode="multiple", activestyle="dotbox",
                       bg="#FFF", fg=Theme.TEXT, font=_f(9), relief="flat",
                       highlightbackground=Theme.LINE, highlightthickness=1)
    lbx_a.pack(side="left", fill="x", expand=True)

    rKB = ttk.Frame(parent, style="Card.TFrame")
    ttk.Label(rKB, text="文件B主键列", style="Sub.TLabel", width=11, anchor="w").pack(side="top", anchor="w")
    var_key_b = tk.StringVar(value=",".join(app.get("key_b", []) or []))
    e_key_b = ttk.Entry(rKB, textvariable=var_key_b, width=30)
    e_key_b.pack(side="top", fill="x", pady=(0, 2))
    lbx_b = tk.Listbox(rKB, width=30, height=5, exportselection=False,
                       selectmode="multiple", activestyle="dotbox",
                       bg="#FFF", fg=Theme.TEXT, font=_f(9), relief="flat",
                       highlightbackground=Theme.LINE, highlightthickness=1)
    lbx_b.pack(side="left", fill="x", expand=True)

    # 数量列（可选）
    rQA = ttk.Frame(parent, style="Card.TFrame")
    ttk.Label(rQA, text="文件A数量列", style="Sub.TLabel", width=11, anchor="w").pack(side="left")
    var_qty_a = tk.StringVar(value=app.get("qty_a", ""))
    cmb_qty_a = ttk.Combobox(rQA, textvariable=var_qty_a, width=28)
    cmb_qty_a.pack(side="left", fill="x", expand=True)

    rQB = ttk.Frame(parent, style="Card.TFrame")
    ttk.Label(rQB, text="文件B数量列", style="Sub.TLabel", width=11, anchor="w").pack(side="left")
    var_qty_b = tk.StringVar(value=app.get("qty_b", ""))
    cmb_qty_b = ttk.Combobox(rQB, textvariable=var_qty_b, width=28)
    cmb_qty_b.pack(side="left", fill="x", expand=True)

    # 核对列：可编辑，默认使用当前 Excel 的位号列
    r1 = ttk.Frame(parent, style="Card.TFrame")
    ttk.Label(r1, text="核对列", style="Sub.TLabel", width=11, anchor="w").pack(side="left")
    cmb_col = ttk.Combobox(r1, textvariable=lb_col, width=28)
    cmb_col.pack(side="left", fill="x", expand=True)

    # 匹配方式
    r2 = ttk.Frame(parent, style="Card.TFrame")
    ttk.Label(r2, text="匹配方式", style="Sub.TLabel", width=11, anchor="w").pack(side="left")
    cmb_mode = ttk.Combobox(r2, textvariable=lb_mode, width=28, state="readonly")
    cmb_mode.pack(side="left", fill="x", expand=True)
    cmb_mode["values"] = ["位号核对(推荐)", "全文匹配", "位号对比"]
    if not lb_mode.get():
        cmb_mode.current(0)

    # 说明
    rflags = ttk.Frame(parent, style="Card.TFrame")
    ttk.Label(rflags,
              text="主键列可多选（Ctrl/Shift 点选多列，或用逗号分隔填写）；\n"
                   "不选时自动识别 MPN/料号/零件号列。主键匹配成功后\n"
                   "再比对数量、位号、厂商等信息。",
              style="Sub.TLabel", wraplength=390, justify="left").pack(side="left", anchor="w")

    def _fill_key_columns(*_):
        """把文件A/B 的表头分别填入对应的多选框，默认选中 MPN/料号 列。"""
        ha = app.get("headers") or []
        hb = app.get("headers_b") or []
        lbx_a.delete(0, "end")
        lbx_b.delete(0, "end")
        for h in ha:
            lbx_a.insert("end", h)
        for h in hb:
            lbx_b.insert("end", h)
        cmb_qty_a["values"] = [""] + list(ha)
        cmb_qty_b["values"] = [""] + list(hb)

        # 同步文本框 <- 列表多选
        def _sync_a(*_):
            sel = [lbx_a.get(i) for i in lbx_a.curselection()]
            if sel:
                var_key_a.set(",".join(sel))
        def _sync_b(*_):
            sel = [lbx_b.get(i) for i in lbx_b.curselection()]
            if sel:
                var_key_b.set(",".join(sel))
        lbx_a.bind("<<ListboxSelect>>", _sync_a)
        lbx_b.bind("<<ListboxSelect>>", _sync_b)

        # 当前已选（文本框中的列名），在列表框中高亮；未指定时默认 MPN/料号
        cur_a = [x.strip() for x in var_key_a.get().split(",") if x.strip() and x.strip() in ha]
        if not cur_a:
            cur_a = [_find_col(ha, _MPN_COL_KEYS)] if ha else []
        for c in cur_a:
            if c in ha:
                lbx_a.selection_set(ha.index(c))
        cur_b = [x.strip() for x in var_key_b.get().split(",") if x.strip() and x.strip() in hb]
        if not cur_b:
            cur_b = [_find_col(hb, _MPN_COL_KEYS)] if hb else []
        for c in cur_b:
            if c in hb:
                lbx_b.selection_set(hb.index(c))
        if cur_a:
            var_key_a.set(",".join(cur_a))
        if cur_b:
            var_key_b.set(",".join(cur_b))
        # 数量列默认
        if ha and (not var_qty_a.get() or var_qty_a.get() not in ha):
            var_qty_a.set(_find_col(ha, _QTY_COL_KEYS) or "")
        if hb and (not var_qty_b.get() or var_qty_b.get() not in hb):
            var_qty_b.set(_find_col(hb, _QTY_COL_KEYS) or "")

    def _fill_columns(*_):
        """把文件A的列名填进核对列下拉。"""
        headers = app.get("headers") or []
        opts = list(headers)
        cmb_col["values"] = opts
        if opts:
            cur = lb_col.get()
            if cur not in opts:
                lb_col.set(_designator_col_name(opts) or opts[0])

    _fill_key_columns()
    _fill_columns()
    cbs = app.get("_col_refresh_callbacks") or []
    if not isinstance(cbs, list):
        cbs = []
    cbs.append(_fill_key_columns)
    cbs.append(_fill_columns)
    app["_col_refresh_callbacks"] = cbs
    lb_xlsx.trace("w", _fill_key_columns)
    lb_xlsx.trace("w", _fill_columns)
    lb_compare.trace("w", _fill_key_columns)

    # 高级设置列显隐：展开/收起
    key_uv = [rKA, rKB, rQA, rQB]

    def _repack_adv(*_):
        adv_rows = [rKA, rKB, rQA, rQB, r1, r2, rflags]
        if not var_adv.get():
            for w in adv_rows:
                w.pack_forget()
            return
        for w in adv_rows:
            w.pack(in_=parent, fill="x", pady=3, before=r3 if r3.winfo_ismapped() else None)
    var_adv.trace("w", _repack_adv)

    # 默认值
    app["out_dir"] = var_out.get()
    app["adv_settings"] = var_adv.get()
    app["key_a"] = var_key_a.get()
    app["key_b"] = var_key_b.get()
    app["qty_a"] = var_qty_a.get()
    app["qty_b"] = var_qty_b.get()
    app["_key_vars"] = (var_key_a, var_key_b, var_qty_a, var_qty_b)

    r3 = ttk.Frame(parent, style="Card.TFrame")
    r3.pack(fill="x", pady=(8, 2))
    ttk.Button(r3, text=" 开始核对 ", style="Primary.TButton", command=lambda: threading.Thread(
        target=_do_check, args=(app, lb_xlsx, lb_pdf, lb_file_b, lb_compare, lb_col, lb_mode,
                                var_status, root),
        daemon=True).start()).pack(side="left", fill="x", expand=True)
    ttk.Button(r3, text="导出报告", style="Success.TButton", command=lambda: threading.Thread(
        target=_do_export, args=(app, root, var_status), daemon=True).start())\
        .pack(side="left", padx=(6, 0), fill="x", expand=True)
    # 初始按当前对比类型与开关同步一次（恢复配置时不会触发 trace）
    _repack_adv()


def _pick_out_dir(app, var_out, root):
    """选择报告输出文件夹，并记住到配置。"""
    d = filedialog.askdirectory(title="选择报告输出文件夹",
                                initialdir=var_out.get() or _smart_dir(app))
    if d:
        var_out.set(d)
        app["out_dir"] = d
        # 立即存配置，防止遗漏
        cfg = _load_config()
        cfg["out_dir"] = d
        _save_config(cfg)


# ---------- 二级窗口按钮 ----------
def _subwin_rows(parent, root, app):
    btns = [
        ("PDF 文字查看", lambda: open_pdf_text_win(root, app)),
        ("不一致清单", lambda: open_mismatch_win(root, app)),
        ("汇总统计", lambda: open_summary_win(root, app)),
    ]
    for text, cmd in btns:
        ttk.Button(parent, text=text, style="TButton", command=cmd).pack(fill="x", pady=2)


# ---------- 预览卡片 ----------
def _preview_card(parent, app, lb_col):
    card = _make_card(parent, "BOM 数据预览")
    card.grid(row=0, column=0, sticky="nsew", pady=(0, 8))
    frame = ttk.Frame(card, style="Card.TFrame")
    frame.pack(fill="both", expand=True)
    tree = ttk.Treeview(frame, show="headings", height=6)
    vs = ttk.Scrollbar(frame, orient="vertical", command=tree.yview)
    hs = ttk.Scrollbar(frame, orient="horizontal", command=tree.xview)
    tree.configure(yscrollcommand=vs.set, xscrollcommand=hs.set)
    tree.pack(side="left", fill="both", expand=True)
    vs.pack(side="right", fill="y")
    hs.pack(side="bottom", fill="x")
    app["preview_tree"] = tree


def _result_card(parent, root, app):
    card = _make_card(parent, "核对结果  —  双击行查看位号详情")
    card.grid(row=1, column=0, sticky="nsew")
    frame = ttk.Frame(card, style="Card.TFrame")
    frame.pack(fill="both", expand=True)
    tree = ttk.Treeview(frame, columns=("row", "col", "des", "found", "pages", "raw"),
                        show="headings", height=18)
    for c, (t, w) in {
        "row": ("BOM行", 60), "col": ("列", 60), "des": ("位号/值", 110),
        "found": ("结果", 74), "pages": ("PDF页码", 90), "raw": ("原值(BOM)", 260),
    }.items():
        tree.heading(c, text=t)
        tree.column(c, width=w, anchor="w")
    vs = ttk.Scrollbar(frame, orient="vertical", command=tree.yview)
    hs = ttk.Scrollbar(frame, orient="horizontal", command=tree.xview)
    tree.configure(yscrollcommand=vs.set, xscrollcommand=hs.set)
    tree.pack(side="left", fill="both", expand=True)
    vs.pack(side="right", fill="y")
    hs.pack(side="bottom", fill="x")
    tree.tag_configure("ok", background=Theme.OK_BG, foreground=Theme.OK_FG)
    tree.tag_configure("bad", background=Theme.BAD_BG, foreground=Theme.BAD_FG)
    tree.tag_configure("odd", background="#FAFCFF")

    def on_double(event):
        # 对比模式下列结构不同，无需详情窗口
        cols = tree["columns"]
        if "des" not in cols:
            return
        sel = tree.selection()
        if not sel:
            return
        vals = tree.item(sel[0], "values")
        if not vals or len(vals) < 6:
            return
        row, col, des, found, pages, raw = vals
        pg = [int(x) for x in str(pages).split(",") if x.strip().isdigit()] if pages else []
        open_detail_win(root, app, des, pg, row, col, raw)

    tree.bind("<Double-1>", on_double)
    app["result_tree"] = tree


# ---------- 核对线程 ----------
def _start_progress(app, root, text):
    """在 GUI 线程启动进度条并更新状态文字。"""
    def _run():
        app.get("progressbar") and app["progressbar"].start(12)
        root.title(f"⏳ {text} — BOM ↔ PDF 器件核对工具 V1")
    root.after(0, _run)


def _stop_progress(app, root, text):
    """停止进度条并恢复标题/状态。"""
    def _run():
        app.get("progressbar") and app["progressbar"].stop()
        root.title("BOM ↔ PDF 器件核对工具 V1")
    root.after(0, _run)


def _do_check(app, lb_xlsx, lb_pdf, lb_file_b, lb_compare, lb_col, lb_mode, var_status, root):
    try:
        ctype = lb_compare.get()
        # 补齐路径
        _sync_paths(app, lb_xlsx, lb_pdf, lb_file_b, lb_compare)

        _start_progress(app, root, "核对中")

        # ===== 方案三：Excel vs Excel（自定义主键列） =====
        if ctype == "Excel vs Excel":
            if not app["bom_path"] or not app["file_b"]:
                _stop_progress(app, root, "")
                root.after(0, lambda: messagebox.showwarning("提示", "请选择 文件A 与 文件B (两个Excel类文件)", parent=root))
                return
            # 读取用户选定的主键/数量列（主键可多列，逗号分隔）
            kv = app.get("_key_vars") or (None, None, None, None)
            ka = [x.strip() for x in kv[0].get().split(",") if x.strip()] if kv[0] else []
            kb = [x.strip() for x in kv[1].get().split(",") if x.strip()] if kv[1] else []
            qa = kv[2].get() if kv[2] else ""
            qb = kv[3].get() if kv[3] else ""
            root.after(0, lambda: var_status.set(
                "正在以主键列 A=%s / B=%s 对比两个 Excel 类文件..." % (",".join(ka) or "自动", ",".join(kb) or "自动")))
            try:
                rows, stats = compare_excel_excel_mpn(app["bom_path"], app["file_b"],
                                                      ka or None, kb or None, qa or None, qb or None)
            except Exception as e:
                _stop_progress(app, root, "")
                err = str(e)
                root.after(0, lambda: messagebox.showerror("错误", "Excel 对比失败:\n%s" % err, parent=root))
                return
            app["result_rows"] = rows
            app["stats"] = stats
            oa, ob = len(stats["only_a"]), len(stats["only_b"])
            _stop_progress(app, root, "完成")
            root.after(0, lambda: _fill_compare_results(app, root, var_status, oa, ob))
            return

        # ===== 方案一：PDF vs PDF =====
        if ctype == "PDF vs PDF":
            if not app["pdf_path"] or not app["file_b"]:
                _stop_progress(app, root, "")
                root.after(0, lambda: messagebox.showwarning("提示", "请选择 文件A 与 文件B (两个PDF)", parent=root))
                return
            root.after(0, lambda: var_status.set("正在解码两个 PDF..."))
            rows, stats = compare_pdf_pdf(app["pdf_path"], app["file_b"])
            app["result_rows"] = rows
            app["stats"] = stats
            oa, ob = len(stats["only_a"]), len(stats["only_b"])
            _stop_progress(app, root, "完成")
            root.after(0, lambda: _fill_compare_results(app, root, var_status, oa, ob))
            return

        # ===== 方案二：Excel vs PDF =====
        if not app["bom_path"]:
            if lb_xlsx.get():
                app["bom_path"] = lb_xlsx.get()
            else:
                _stop_progress(app, root, "")
                root.after(0, lambda: messagebox.showwarning("提示", "请选择 Excel BOM 文件"))
                return
        if not app["pdf_path"]:
            if lb_pdf.get():
                app["pdf_path"] = lb_pdf.get()
            else:
                _stop_progress(app, root, "")
                root.after(0, lambda: messagebox.showwarning("提示", "请选择 PDF 文件"))
                return

        root.after(0, lambda: var_status.set("正在解码 PDF 原理图文字..."))

        # 预热：若未载入 BOM
        if not app["headers"]:
            try:
                headers, data = load_bom(app["bom_path"])
                app["headers"], app["data"] = headers, data
            except Exception as e:
                err = str(e)
                _stop_progress(app, root, "")
                root.after(0, lambda: messagebox.showerror("错误", "BOM 读取失败:\n%s" % err))
                return

        pages_words, pages_text, _ = decode_pdf_blocks(app["pdf_path"])
        app["pdf_pages_text"] = pages_text
        pdf_tokens, all_words = extract_designators(pages_words)
        app["pdf_tokens"] = pdf_tokens
        app["pdf_all_words"] = all_words

        # 位号上下文（供详情窗口）
        ctx = {}
        for pno, pwords in pages_words.items():
            for y, words in pwords:
                for w, x0, x1, top in words:
                    up = w.upper()
                    if re.match(r"^[A-Za-z]{1,4}\d{1,4}$", up):
                        ctx.setdefault(up, []).append(f"第{pno}页  y≈{y:.0f}   {w}")
        app["designator_context"] = ctx

        col_idx = app["headers"].index(lb_col.get()) if lb_col.get() in app["headers"] else 0
        mode = lb_mode.get()
        if "位号" in mode and "对比" not in mode:
            rows = check_designators(app["headers"], app["data"], col_idx, pdf_tokens)
        else:
            rows = check_value_in_pdf(app["headers"], app["data"], col_idx, all_words,
                                      designators_only=False)

        total, okc, badc, bad_list = CompareResult.stat(rows)
        all_bom = set()
        for r in rows:
            for d in r["designators"]:
                all_bom.add(d["designator"])
        extra = sorted(set(pdf_tokens.keys()) - all_bom)
        stats = {
            "all": {
                "BOM位号总数": total,
                "一致(在PDF中找到)": okc,
                "不一致(未找到)": badc,
                "PDF识别位号总数": len(pdf_tokens),
                "PDF中有而BOM无": len(extra),
            },
            "bad": [dict(b, row=r["header_row"]) for r in rows for b in r["designators"]
                    if not b["found"]],
            "extra": extra,
            "in_bom": all_bom,
        }
        app["result_rows"] = rows
        app["stats"] = stats

        _stop_progress(app, root, "完成")
        root.after(0, lambda: _fill_results(app, root, var_status, total, okc, badc, len(extra)))
    except Exception as e:
        import traceback
        traceback.print_exc()
        err = str(e)
        _stop_progress(app, root, "")
        root.after(0, lambda: messagebox.showerror("错误", "核对失败:\n%s" % err))


def _sync_paths(app, lb_xlsx, lb_pdf, lb_file_b, lb_compare):
    """把界面输入同步进 app。"""
    if lb_xlsx.get() and not app["bom_path"]:
        app["bom_path"] = lb_xlsx.get()
    if lb_pdf.get() and not app["pdf_path"]:
        app["pdf_path"] = lb_pdf.get()
    if lb_file_b.get() and not app["file_b"]:
        app["file_b"] = lb_file_b.get()


def _fill_compare_results(app, root, var_status, oa, ob):
    tree = app.get("result_tree")
    if not tree:
        return
    for i in tree.get_children():
        tree.delete(i)
    mode = app.get("stats", {}).get("mode", "set")
    if mode == "mpn":
        tree["columns"] = ("item", "status", "qty_a", "qty_b", "refs_a", "refs_b", "mfg", "pages")
        for c, (t, w) in {"item": ("MPN(料号)", 150), "status": ("核对结果", 170),
                          "qty_a": ("Excel数量", 70), "qty_b": ("PDF数量", 70),
                          "refs_a": ("Excel位号", 220), "refs_b": ("PDF位号", 220),
                          "mfg": ("厂商", 100), "pages": ("PDF页", 60)}.items():
            tree.heading(c, text=t)
            tree.column(c, width=w, anchor="w")
        for row in app["result_rows"]:
            st = row["status"]
            tag = "ok" if st == "一致" else "bad"
            tree.insert("", "end",
                        values=(row["item"], st, row["qty_a"], row["qty_b"], row["refs_a"],
                                row["refs_b"], row["mfg"], row["pages"]),
                        tags=(tag,))
        al = app.get("stats", {}).get("all", {})
        n_ok = al.get("一致(数量/位号匹配)", 0)
        n_bad = al.get("待确认(数量或位号提示)", 0)
        n_only = al.get("PDF有而Excel无", 0)
        n_only_a = al.get("仅Excel有(PDF无此料号)", 0)
        var_status.set(f"MPN 对比完成：一致 {n_ok}，待确认 {n_bad}，仅Excel有 {n_only_a}")
        messagebox.showinfo("MPN 对比完成",
                            f"Excel 料号总数: {al.get('Excel 料号(MPN)总数', 0)}\n"
                            f"一致(数量/位号匹配): {n_ok}\n待确认: {n_bad}\n"
                            f"仅Excel有(PDF无此料号): {n_only_a}\nPDF有而Excel无: {n_only}",
                            parent=root)
        return
    if mode == "excelmpn":
        # Excel vs Excel：MPN 主键，比对数量/位号/厂商
        tree["columns"] = ("item", "status", "qtyA", "qtyB", "refA", "refB", "mfgA", "mfgB")
        for c, (t, w) in {"item": ("MPN(物料号)", 170), "status": ("结果", 120),
                          "qtyA": ("数量A", 60), "qtyB": ("数量B", 60),
                          "refA": ("位号A", 200), "refB": ("位号B", 200),
                          "mfgA": ("厂商A", 90), "mfgB": ("厂商B", 90)}.items():
            tree.heading(c, text=t)
            tree.column(c, width=w, anchor="w")
        for row in app["result_rows"]:
            st = row["status"]
            tag = "ok" if st == "一致" else "bad"
            ra = row.get("rec_a") or {}
            rb = row.get("rec_b") or {}
            tree.insert("", "end",
                        values=(row["item"], st,
                                ra.get("qty", ""), rb.get("qty", ""),
                                ",".join(ra.get("refs", [])), ",".join(rb.get("refs", [])),
                                ra.get("mfg", ""), rb.get("mfg", "")),
                        tags=(tag,))
        al = app.get("stats", {}).get("all", {})
        n_ok = al.get("一致(数量/位号/厂商匹配)", 0)
        n_diff = al.get("字段不一致", 0)
        var_status.set(f"Excel 对比完成：一致 {n_ok}，字段不一致 {n_diff}，仅A有 {oa}，仅B有 {ob}")
        messagebox.showinfo("Excel MPN 对比完成",
                            f"MPN 总数: {al.get('MPN 总数', 0)}\n一致: {n_ok}\n字段不一致: {n_diff}\n"
                            f"仅文件A有: {oa}\n仅文件B有: {ob}",
                            parent=root)
        return
    if mode == "pdf2excel":
        tree["columns"] = ("item", "status", "valueA", "footA", "near")
        for c, (t, w) in {"item": ("PDF位号", 100), "status": ("核对结果", 130),
                          "valueA": ("Excel值(Value)", 130), "footA": ("Excel封装(Footprint)", 160),
                          "near": ("PDF附近标注", 340)}.items():
            tree.heading(c, text=t)
            tree.column(c, width=w, anchor="w")
        for row in app["result_rows"]:
            st = row["status"]
            tag = "ok" if st == "一致" else "bad"
            tree.insert("", "end",
                        values=(row["item"], st, row["valueA"], row["footA"], row["near"]),
                        tags=(tag,))
        n_ok = app.get("stats", {}).get("all", {}).get("值疑似一致", 0)
        n_may = app.get("stats", {}).get("all", {}).get("值疑似不一致/待确认", 0)
        n_miss = app.get("stats", {}).get("all", {}).get("PDF有Excel无", 0)
        var_status.set(f"器件核对完成：值一致 {n_ok}，待确认 {n_may}，PDF有Excel无 {n_miss}")
        messagebox.showinfo("器件核对完成",
                            f"值疑似一致: {n_ok}\n值疑似不一致/待确认: {n_may}\nPDF有Excel无: {n_miss}",
                            parent=root)
        return
    else:
        tree["columns"] = ("item", "status", "a", "b")
        for c, (t, w) in {"item": ("器件/位号", 130), "status": ("结果", 90),
                          "a": ("文件A", 150), "b": ("文件B", 150)}.items():
            tree.heading(c, text=t)
            tree.column(c, width=w, anchor="w")
        for row in app["result_rows"]:
            st = row["status"]
            tag = "ok" if st == "一致" else "bad"
            tree.insert("", "end", values=(row["item"], st, row["a_pages"] or row["a"],
                                           row["b_pages"] or row["b"]), tags=(tag,))
    n_common = len(app.get("stats", {}).get("common", []))
    var_status.set(f"对比完成：一致 {n_common}，仅A有 {oa}，仅B有 {ob}")
    total = len(app["result_rows"])
    messagebox.showinfo("对比完成",
                        f"共 {total} 项\n 一致: {n_common}\n 仅A有: {oa}  仅B有: {ob}",
                        parent=root)


def _fill_results(app, root, var_status, total, okc, badc, nextra):
    tree = app.get("result_tree")
    if not tree:
        return
    for i in tree.get_children():
        tree.delete(i)
    odd = False
    for row in app["result_rows"]:
        for d in row["designators"]:
            ok = d["found"]
            odd = not odd
            tree.insert("", "end",
                        values=(row["header_row"], row["col_name"], d["designator"],
                                "一致" if ok else "不一致",
                                ",".join(str(p) for p in d["pages"]) if ok else "",
                                str(row["raw"])[:90]),
                        tags=("ok" if ok else "bad",)) if ok else tree.insert(
                "", "end",
                values=(row["header_row"], row["col_name"], d["designator"], "不一致",
                        "", str(row["raw"])[:90]),
                tags=("bad",))
    var_status.set(f"核对完成：共 {total} 个位号，一致 {okc}，不一致 {badc}，PDF额外位号 {nextra}")
    messagebox.showinfo("核对完成",
                        f"共 {total} 个位号\n一致: {okc}\n不一致: {badc}\nPDF额外位号: {nextra}",
                        parent=root)


# ---------- 导出报告 ----------
def _do_export(app, root, var_status):
    try:
        if not app.get("stats"):
            root.after(0, lambda: messagebox.showwarning("提示", "请先执行“开始核对”。", parent=root))
            return
        _start_progress(app, root, "导出报告中")
        d = app.get("out_dir") or os.path.dirname(os.path.abspath(__file__))
        fmt_vars = app.get("_fmt_vars")
        if fmt_vars:
            fmts = set()
            for key, var in zip(("xlsx", "docx", "pdf", "txt"), fmt_vars):
                if var.get():
                    fmts.add(key)
        else:
            fmts = {"xlsx", "docx", "pdf", "txt"}
        xlsx_path, txt_path, docx_path = gen_report(app["bom_path"], app["result_rows"], app["stats"],
                                                    app["pdf_tokens"], app["pdf_pages_text"], d, fmts)
        _stop_progress(app, root, "完成")
        # 生成消息（按选中类型罗列实际产出文件）
        made = []
        for var, p_ in (("xlsx", xlsx_path), ("txt", txt_path), ("docx", docx_path)):
            if var in fmts and p_:
                made.append(p_)
        if "pdf" in fmts:
            import glob as _g
            base_pdf = (os.path.splitext(os.path.basename(xlsx_path or txt_path or docx_path or "r"))[0]
                        + ".pdf")
            cand = os.path.join(d, base_pdf)
            if os.path.exists(cand):
                made.append(cand)
        msg = "报告已生成:\n" + "\n".join(made)
        app["last_report_msg"] = msg
        root.after(0, lambda: (var_status.set(msg), messagebox.showinfo("完成", msg, parent=root)))
    except Exception as e:
        traceback.print_exc()
        err = str(e)
        _stop_progress(app, root, "")
        root.after(0, lambda: messagebox.showerror("错误", "报告生成失败:\n%s" % err, parent=root))


# ---------- 自动载入 ----------
def _auto_load(app, lb_xlsx, lb_pdf, lb_col):
    import glob
    d = _smart_dir(app)
    xlsx = (glob.glob(os.path.join(d, "*.BOM.xlsx")) + glob.glob(os.path.join(d, "*BOM*.xlsx"))
            + glob.glob(os.path.join(d, "*.BOM.csv")) + glob.glob(os.path.join(d, "*BOM*.csv"))
            + glob.glob(os.path.join(d, "*.BOM.tsv")) + glob.glob(os.path.join(d, "*BOM*.tsv"))
            + glob.glob(os.path.join(d, "*.xlsx")) + glob.glob(os.path.join(d, "*.csv"))
            + glob.glob(os.path.join(d, "*.tsv")))
    pdf = glob.glob(os.path.join(d, "*.pdf"))
    if not app.get("bom_path") and xlsx:
        p = sorted(xlsx)[0]
        app["bom_path"] = p
        lb_xlsx.set(os.path.basename(p))
        try:
            headers, data = load_bom(p)
            app["headers"], app["data"] = headers, data
        except Exception:
            pass
    if not app.get("pdf_path") and pdf:
        p = sorted(pdf)[0]
        app["pdf_path"] = p
        lb_pdf.set(os.path.basename(p))
    _refresh_preview(app, lb_col)


# ============================================================
#  命令行入口（GUI / 命令行核对 / 帮助 三模式）
# ============================================================
def run_cli(bom_path=None, pdf_path=None, col_name="Part Reference", mode="位号核对(推荐)",
            out_dir=None, file_b=None, ctype=None, formats=None):
    """命令行核对/对比模式：无需图形界面即可输出结果并生成报告。
    ctype: Excel vs Excel / Excel vs PDF / PDF vs PDF
    formats: {"xlsx","docx","pdf","txt"} 子集，None=全部
    """
    if formats is None:
        formats = {"xlsx", "docx", "pdf", "txt"}
    else:
        formats = set(formats)

    def _print_paths(xlsx_path, txt_path, docx_path):
        print("报告已生成:")
        if "xlsx" in formats and xlsx_path:
            print("   %s" % xlsx_path)
        if "txt" in formats and txt_path:
            print("   %s" % txt_path)
        if "docx" in formats and docx_path:
            print("   %s" % docx_path)
        if "pdf" in formats:
            base = os.path.splitext(os.path.basename(xlsx_path or txt_path or docx_path))[0]
            p = os.path.join(out_dir, base + ".pdf") if out_dir and base else None
            if p and os.path.exists(p):
                print("   %s" % p)

    # 方案三：Excel vs Excel（MPN 主键）
    if ctype == "Excel vs Excel":
        if not bom_path or not file_b:
            print("用法: --check --compare 'Excel vs Excel' --a <文件A> --b <文件B>")
            return 1
        print("[1/3] 以主键列对比两个 Excel 类文件...")
        try:
            rows, stats = compare_excel_excel_mpn(bom_path, file_b, col_name or None)
        except Exception as e:
            print("错误: %s" % e)
            return 1
        oa, ob = len(stats["only_a"]), len(stats["only_b"])
        print("[2/3] 结果: 一致 %d  字段不一致 %d  仅A有 %d  仅B有 %d"
              % (stats["all"]["一致(数量/位号/厂商匹配)"], stats["all"]["字段不一致"], oa, ob))
        print("[3/3] 生成报告...")
        if not out_dir:
            out_dir = os.path.dirname(os.path.abspath(bom_path))
        xlsx_path, txt_path, docx_path = gen_report(bom_path, rows, stats, {}, {}, out_dir, formats)
        _print_paths(xlsx_path, txt_path, docx_path)
        return 0 if oa == 0 and ob == 0 else 2

    # 方案一：PDF vs PDF
    if ctype == "PDF vs PDF":
        if not bom_path or not file_b:
            print("用法: --check --compare 'PDF vs PDF' --a <文件A> --b <文件B>")
            return 1
        print("[1/3] 解码并对比两个 PDF...")
        rows, stats = compare_pdf_pdf(bom_path, file_b)
        oa, ob = len(stats["only_a"]), len(stats["only_b"])
        print("[2/3] 结果: 两文件一致 %d  仅A有 %d  仅B有 %d" % (len(stats["common"]), oa, ob))
        print("[3/3] 生成报告...")
        if not out_dir:
            out_dir = os.path.dirname(os.path.abspath(bom_path))
        xlsx_path, txt_path, docx_path = gen_report(bom_path, rows, stats, {}, {}, out_dir, formats)
        _print_paths(xlsx_path, txt_path, docx_path)
        return 0 if oa == 0 and ob == 0 else 2

    # 方案二：Excel vs PDF（BOM 器件核对）
    if not bom_path or not pdf_path:
        print("用法: bom_pdf_verify.py --check --bom <BOM.xlsx> --pdf <原理图.pdf> [--col 列名]")
        print("      或双击运行进入 GUI 界面")
        return 1
    headers, data = load_bom(bom_path)
    if col_name not in headers:
        # 自动识别位号列（Reference / Part Reference / 位号…）
        auto = _designator_col_name(headers) or (headers[0] if headers else None)
        if auto and (col_name == "Part Reference" or not col_name):
            col_name = auto
        if col_name not in headers:
            print("错误: 列名 '%s' 不存在。可用列: %s" % (col_name, ", ".join(headers)))
            return 1
    print("[1/4] 解码 PDF 原理图文字...")
    pages_words, pages_text, _ = decode_pdf_blocks(pdf_path)
    pdf_tokens, all_words = extract_designators(pages_words)
    col_idx = headers.index(col_name)
    if "位号" in mode:
        rows = check_designators(headers, data, col_idx, pdf_tokens)
    else:
        rows = check_value_in_pdf(headers, data, col_idx, all_words, designators_only=False)
    total, okc, badc, bad_list = CompareResult.stat(rows)
    all_bom = set()
    for r in rows:
        for d in r["designators"]:
            all_bom.add(d["designator"])
    extra = sorted(set(pdf_tokens.keys()) - all_bom)
    stats = {
        "all": {
            "BOM位号总数": total,
            "一致(在PDF中找到)": okc,
            "不一致(未找到)": badc,
            "PDF识别位号总数": len(pdf_tokens),
            "PDF中有而BOM无": len(extra),
        },
        "bad": [dict(b, row=r["header_row"]) for r in rows for b in r["designators"]
                if not b["found"]],
        "extra": extra,
        "in_bom": all_bom,
    }
    print("[2/4] 核对结果: 共 %d 个位号，一致 %d，不一致 %d" % (total, okc, badc))
    print("[3/4] 不一致(未在PDF中找到): %d" % len(bad_list))
    for b in bad_list:
        print("      MISS  %s  (BOM第%d行)" % (b["designator"], 0))
    print("[4/4] 生成报告...")
    if not out_dir:
        out_dir = os.path.dirname(os.path.abspath(bom_path))
    xlsx_path, txt_path, docx_path = gen_report(bom_path, rows, stats, pdf_tokens, pages_text, out_dir, formats)
    _print_paths(xlsx_path, txt_path, docx_path)
    return 0 if badc == 0 else 2


def main():
    import argparse
    parser = argparse.ArgumentParser(
        description="BOM(Excel) 与 PDF 原理图 器件核对 / 对比工具",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="示例:\n"
               "  双击 / 无参数      : 打开 GUI 界面（支持三种对比方案）\n"
               "  方案一 PDF-PDF    : --check --compare 'PDF vs PDF' --a a.pdf --b b.pdf\n"
               "  方案二 Excel-PDF  : --check --bom xxx.xlsx --pdf xxx.pdf\n"
               "  方案三 Excel-Excel: --check --compare 'Excel vs Excel' --a a.csv --b b.xlsx\n"
               "  指定输出目录      : 加 --out ./report\n")
    parser.add_argument("--gui", action="store_true", help="启动图形界面(默认)")
    parser.add_argument("--check", action="store_true", help="命令行模式(无需GUI)")
    parser.add_argument("--compare", type=str, default=None,
                        choices=["Excel vs Excel", "PDF vs PDF"], help="双文件对比模式")
    parser.add_argument("--a", type=str, default=None, help="文件A 路径")
    parser.add_argument("--b", type=str, default=None, help="文件B 路径(对比用)")
    parser.add_argument("--bom", type=str, default=None, help="BOM Excel (.xlsx) 路径")
    parser.add_argument("--pdf", type=str, default=None, help="原理图 PDF 路径")
    parser.add_argument("--col", type=str, default="Part Reference", help="要核对的列名")
    parser.add_argument("--match", type=str, default="位号核对(推荐)",
                        choices=["位号核对(推荐)", "全文匹配"], help="匹配方式")
    parser.add_argument("--out", type=str, default=None, help="报告输出目录(默认=A所在目录)")
    parser.add_argument("--format", type=str, default=None,
                        help="导出类型，逗号分隔，如 xlsx,docx,pdf,txt（默认全部）")
    args = parser.parse_args()

    fmts = None
    if args.format:
        alias = {"excel": "xlsx", "word": "docx", "pdf": "pdf", "txt": "txt",
                 "xlsx": "xlsx", "docx": "docx"}
        fmts = list({alias.get(x.strip().lower(), x.strip().lower())
                     for x in args.format.split(",") if x.strip()})

    if args.compare:
        code = run_cli(args.a, args.pdf, args.col, args.match, args.out, args.b, args.compare, fmts)
        sys.exit(code)
    if args.check:
        code = run_cli(args.bom, args.pdf, args.col, args.match, args.out, formats=fmts)
        sys.exit(code)
    # 默认进入 GUI
    run_gui()


if __name__ == "__main__":
    main()


